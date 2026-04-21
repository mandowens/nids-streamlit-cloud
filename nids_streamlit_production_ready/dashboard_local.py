from __future__ import annotations

import contextlib
import hashlib
import hmac
import importlib.util
import py_compile
import io
import json
import os
import platform
import sqlite3
import sys
import time
import traceback
import zipfile
import shutil
import signal
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
import copy

try:
    from streamlit_autorefresh import st_autorefresh
except Exception:  # optional dependency
    def st_autorefresh(*args, **kwargs):
        return None

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # optional dependency
    Document = None
    Inches = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
except Exception:  # optional dependency
    A4 = None
    canvas = None
    ImageReader = None

APP_DIR = Path(__file__).resolve().parent
DEFAULT_PIPELINE_PATH = APP_DIR / "pipeline_local.py"
FALLBACK_PIPELINE_PATH = APP_DIR / "pipeline_local.py"
UPLOAD_ROOT = APP_DIR / "uploaded_datasets"
ARTIFACT_ROOT = APP_DIR / "dashboard_local_full_artifacts"
JOB_ROOT = APP_DIR / "dashboard_local_full_jobs"
DB_PATH = APP_DIR / "nids_dashboard_local_full.sqlite"

DATASET_LABELS = {
    "nslkdd": "NSL-KDD",
    "unsw": "UNSW-NB15",
    "insdn": "InSDN",
    "cicids": "CICIDS2017",
}

DEFAULT_PATHS = {
    "nslkdd": {"train": "data/nslkdd/KDDTrain+.txt", "test": "data/nslkdd/KDDTest+.txt"},
    "unsw": {"train": "data/unsw/UNSW_NB15_training-set.csv", "test": "data/unsw/UNSW_NB15_testing-set.csv"},
    "insdn": {"normal": "data/insdn/Normal_data.csv", "ovs": "data/insdn/OVS_data.csv", "metasploit": "data/insdn/Metasploit.csv"},
    "cicids": {"dir": "data/cicids2017"},
}

SCALER_LABELS = {
    "standard": "StandardScaler",
    "minmax": "MinMaxScaler",
    "robust": "RobustScaler",
}
BALANCE_LABELS = {
    "smote": "SMOTE",
    "adasyn": "ADASYN",
    "borderline": "BorderlineSMOTE",
    "under": "RandomUnderSampler",
}
COPULA_LABELS = {
    "gaussian": "Gaussian Copula",
    "clayton": "Clayton Copula",
    "frank": "Frank Copula",
    "gumbel": "Gumbel Copula",
    "student_t": "Student-t Copula",
    "vine": "Vine Copula (pair-copula approximation)",
}

VALIDATION_METHOD_LABELS = {
    "holdout": "Hold-Out",
    "simple_kfold": "Simple K-Fold",
    "stratified_kfold": "Stratified K-Fold",
    "repeated_kfold": "Repeated K-Fold",
    "bootstrapping": "Bootstrapping",
    "nested_cv": "Nested Cross-Validation",
}
FEATURE_METHOD_LABELS = {
    "mutual_info": "Mutual Information",
    "anova": "ANOVA F-test",
    "chi2": "Chi-Square",
    "l1_logistic": "L1 Logistic",
    "rf_importance": "Random Forest Importance",
    "rfe": "Recursive Feature Elimination",
    "mrmr": "mRMR approximation",
}
ENSEMBLE_METHOD_LABELS = {
    "soft_vote": "Soft Vote",
    "weighted_soft_vote": "Weighted Soft Vote",
    "hard_vote": "Hard Vote",
    "geometric_mean": "Geometric Mean",
    "rank_average": "Rank Average",
    "stacking": "Stacking",
}
HYBRID_METHOD_LABELS = {
    "weighted_ml_dl": "Weighted ML vs DL",
    "stack_groups": "Stack groups (ML vs DL)",
}

LEARNING_PARADIGM_LABELS = {
    "supervised": "Supervised Learning",
    "semi_supervised": "Semi-Supervised Learning",
    "unsupervised": "Unsupervised Learning",
    "reinforcement": "Reinforcement Learning",
}

SUPERVISED_TRACK_LABELS = {
    "ml": "Supervised → Model ML",
    "dl": "Supervised → Model DL",
    "both": "Supervised → Model ML + Model DL",
    "hybrid": "Supervised → Ensemble / Hybrid / Fusion",
}

SEMI_ML_MODEL_SPECS = [
    ("use_label_spreading", "LabelSpreading", "Graph-based semi-supervised label propagation"),
    ("use_self_training_svm", "SelfTraining-SVM", "Self-training with SVM base estimator"),
    ("use_self_training_rf", "SelfTraining-RF", "Self-training with Random Forest base estimator"),
]
SEMI_DL_MODEL_SPECS = [
    ("use_pseudo_label_dnn", "PseudoLabelDNN", "Semi-supervised pseudo-label dense neural network"),
]

UNSUP_ML_MODEL_SPECS = [
    ("use_isolation_forest", "IsolationForest", "Tree-based anomaly detector"),
    ("use_oneclass_svm_anom", "OneClassSVM", "Boundary-based anomaly detector"),
    ("use_lof", "LOF", "Local Outlier Factor (novelty mode)"),
    ("use_pca_anom", "PCARecon", "PCA reconstruction error anomaly detector"),
    ("use_kmeans_anom", "KMeansDistance", "Distance-to-centroid anomaly detector"),
]
UNSUP_DL_MODEL_SPECS = [
    ("use_autoencoder_anom", "AE-Anomaly", "Autoencoder reconstruction-based anomaly detection"),
]

RL_ML_MODEL_SPECS = [
    ("use_rl_threshold", "AdaptiveQThreshold*", "Experimental Q-learning threshold policy over anomaly scores"),
]
RL_DL_MODEL_SPECS = []

ROADMAP_MODEL_GROUPS = {
    "reinforcement": {
        "DL": ["Deep Q-Network (roadmap)", "PPO (roadmap)", "A2C (roadmap)"],
    },
}

SCALER_RECOMMENDATION = {
    "nslkdd": {"recommended": "robust", "reason": "Campuran fitur count/rate/bytes cukup rentan terhadap outlier dan rentang yang timpang."},
    "unsw": {"recommended": "robust", "reason": "Fitur numerik UNSW-NB15 heterogen dan sebarannya lebar; robust aman sebagai baseline."},
    "insdn": {"recommended": "robust", "reason": "Penggabungan beberapa sumber traffic menghasilkan variasi besar; robust lebih tahan spike/outlier."},
    "cicids": {"recommended": "robust", "reason": "Flow features CICIDS2017 sering ekstrem dan sangat beragam; robust cocok untuk baseline awal."},
}
COPULA_DEFAULT = {"nslkdd": "gaussian", "unsw": "clayton", "insdn": "frank", "cicids": "gaussian"}

SPLIT_PROTOCOLS = {
    "nslkdd": "Official train/test split dipertahankan; validation carved from official train only.",
    "unsw": "Official training/testing split dipertahankan; validation carved from official training set only.",
    "insdn": "Multiple traffic sources are merged first, then stratified into train/validation/test.",
    "cicids": "Multiple daily flow files are merged first, then split with random or temporal protocol depending on configuration.",
}

WORKFLOW_MODES = {
    "exploratory": "Exploratory mode",
    "publication": "Publication mode",
}

EXPERIMENT_TYPES = {
    "single_baseline": "Single baseline",
    "ml_ensemble": "ML ensemble",
    "dl_ensemble": "DL ensemble",
    "hybrid_fusion": "Hybrid ML + DL fusion",
    "copula_bn": "Copula fusion + Bayesian reasoning",
    "custom": "Custom design",
}

PRESET_OPTIONS = ["Fast baseline", "Balanced research", "Full dissertation", "Custom"]
PRESET_DEFAULTS = {
    "Fast baseline": dict(
        cv_folds=2, top_k=30, sample_frac=0.10,
        use_lr=True, use_rf=True, use_svm=False, use_xgb=True, use_dt=True,
        use_mlp=True, use_dnn=False, use_cnn1d=False, use_lstm=False,
        use_aae=False, use_cnn_lstm=False, use_cnn_bilstm=False, use_cnn_gru=False,
        use_gnn=False, use_gat=False, use_transformer=False, use_transfer_learning=False,
        use_bn=False, use_shap=False,
    ),
    "Balanced research": dict(
        cv_folds=2, top_k=40, sample_frac=0.30,
        use_lr=True, use_rf=True, use_svm=True, use_xgb=True, use_dt=True,
        use_mlp=True, use_dnn=True, use_cnn1d=True, use_lstm=True,
        use_aae=False, use_cnn_lstm=True, use_cnn_bilstm=True, use_cnn_gru=True,
        use_gnn=False, use_gat=False, use_transformer=True, use_transfer_learning=False,
        use_bn=True, use_shap=False,
    ),
    "Full dissertation": dict(
        cv_folds=5, top_k=40, sample_frac=1.0,
        use_lr=True, use_rf=True, use_svm=True, use_xgb=True, use_dt=True,
        use_mlp=True, use_dnn=True, use_cnn1d=True, use_lstm=True,
        use_aae=False, use_cnn_lstm=True, use_cnn_bilstm=True, use_cnn_gru=True,
        use_gnn=True, use_gat=True, use_transformer=True, use_transfer_learning=True,
        use_bn=True, use_shap=True,
    ),
    "Custom": dict(
        cv_folds=2, top_k=40, sample_frac=0.30,
        use_lr=True, use_rf=True, use_svm=True, use_xgb=True, use_dt=True,
        use_mlp=True, use_dnn=False, use_cnn1d=False, use_lstm=False,
        use_aae=False, use_cnn_lstm=False, use_cnn_bilstm=True, use_cnn_gru=True,
        use_gnn=False, use_gat=False, use_transformer=True, use_transfer_learning=False,
        use_bn=True, use_shap=False,
    ),
}

DATASET_FINAL_PRESETS = {
    "nslkdd": {
        "None": {"description": "Tanpa overlay preset dataset."},
        "NSL-KDD Final Baseline": {"description": "Robust + SMOTE + Gaussian + stratified k-fold.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "gaussian", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 40, "feature_method": "mutual_info"},
        "NSL-KDD Hybrid Final": {"description": "ML+DL hybrid final preset untuk NSL-KDD.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "gaussian", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 45, "feature_method": "mutual_info", "use_rf": True, "use_xgb": True, "use_svm": True, "use_dnn": True, "use_cnn1d": True, "use_lstm": True, "use_cnn_bilstm": True, "use_transformer": True}
    },
    "unsw": {
        "None": {"description": "Tanpa overlay preset dataset."},
        "UNSW Final Baseline": {"description": "Robust + SMOTE + Clayton + stratified k-fold.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "clayton", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 45, "feature_method": "mrmr"},
        "UNSW Hybrid Final": {"description": "ML+DL hybrid final preset untuk UNSW-NB15.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "clayton", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 50, "feature_method": "mrmr", "use_rf": True, "use_xgb": True, "use_svm": True, "use_dnn": True, "use_lstm": True, "use_cnn_lstm": True, "use_transformer": True}
    },
    "cicids": {
        "None": {"description": "Tanpa overlay preset dataset."},
        "CICIDS Publication": {"description": "Baseline publikasi CICIDS2017 dengan sample 0.30.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "gaussian", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 50, "sample_frac": 0.30, "feature_method": "rf_importance"},
        "CICIDS Full Hybrid": {"description": "Preset hybrid final CICIDS2017 dengan full sample.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "gaussian", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 60, "sample_frac": 1.0, "feature_method": "rf_importance", "use_rf": True, "use_xgb": True, "use_svm": True, "use_dnn": True, "use_cnn1d": True, "use_lstm": True, "use_cnn_bilstm": True, "use_transformer": True}
    },
    "insdn": {
        "None": {"description": "Tanpa overlay preset dataset."},
        "InSDN Final Baseline": {"description": "Robust + SMOTE + Frank untuk InSDN.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "frank", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 50, "feature_method": "rf_importance"},
        "InSDN Hybrid Final": {"description": "Preset hybrid final InSDN dengan kombinasi tabular-ML dan sequence-DL.", "binary": True, "scaler": "robust", "balance": "smote", "copula_family": "frank", "validation_method": "stratified_kfold", "cv_folds": 5, "top_k": 55, "feature_method": "mrmr", "use_rf": True, "use_xgb": True, "use_dnn": True, "use_cnn1d": True, "use_lstm": True, "use_cnn_lstm": True, "use_transformer": True}
    }
}

FINAL_STUDY_PRESETS = {
    "Custom": {"description": "Tanpa preset final study tambahan."},
    "Binary Final Study": {
        "description": "Preset penelitian akhir untuk klasifikasi biner dengan penekanan pada stabilitas dan interpretabilitas.",
        "binary": True,
        "workflow_mode": "publication",
        "experiment_type": "hybrid_fusion",
        "optimize_ml": True,
        "optimize_dl": True,
        "use_ensemble": True,
        "use_hybrid": True,
        "ensemble_method": "weighted_soft_vote",
        "hybrid_method": "weighted_ml_dl",
        "generate_eda": True,
        "use_bn": True,
    },
    "Multiclass Final Study": {
        "description": "Preset penelitian akhir untuk klasifikasi multikelas dengan pelaporan kelas-per-kelas dan evaluasi yang lebih ketat.",
        "binary": False,
        "workflow_mode": "publication",
        "experiment_type": "hybrid_fusion",
        "optimize_ml": True,
        "optimize_dl": True,
        "use_ensemble": True,
        "use_hybrid": True,
        "ensemble_method": "stacking",
        "hybrid_method": "stack_groups",
        "generate_eda": True,
        "use_bn": True,
        "top_k": 50,
    },
}

PUBLICATION_VALIDATION_PRESETS = {
    "Custom": {"description": "Tanpa preset validasi publikasi tambahan."},
    "Publication Hold-Out": {
        "description": "Hold-out final dengan validasi standar dan biaya komputasi lebih ringan.",
        "validation_method": "holdout",
        "cv_folds": 5,
    },
    "Publication Stratified CV": {
        "description": "Stratified K-Fold untuk publikasi baseline yang kuat dan seimbang.",
        "validation_method": "stratified_kfold",
        "cv_folds": 5,
    },
    "Publication Nested CV": {
        "description": "Nested Cross-Validation untuk evaluasi kuat pada tahap seleksi/optimasi model.",
        "validation_method": "nested_cv",
        "cv_folds": 5,
        "nested_inner_folds": 3,
        "optimize_ml": True,
    },
}



THEME_OPTIONS = {
    "cyber": {"label": "Cyber Neon", "bg1": "rgba(55,120,255,0.18)", "bg2": "rgba(0,245,255,0.14)", "base1": "#07111f", "base2": "#0a1425", "base3": "#0c1730", "accent": "#6ee7ff", "accent2": "#7c9dff", "ok": "#5cf0a5", "warn": "#ffd166", "bad": "#ff7b95"},
    "aurora": {"label": "Aurora Research", "bg1": "rgba(34,197,94,0.16)", "bg2": "rgba(56,189,248,0.13)", "base1": "#08131a", "base2": "#0f1e29", "base3": "#132734", "accent": "#67e8f9", "accent2": "#86efac", "ok": "#86efac", "warn": "#fde68a", "bad": "#fca5a5"},
    "academic": {"label": "Academic Clean", "bg1": "rgba(43,96,185,0.10)", "bg2": "rgba(167,189,255,0.10)", "base1": "#f4f7fb", "base2": "#ecf1f8", "base3": "#e4ebf5", "accent": "#1d4ed8", "accent2": "#5b7cfa", "ok": "#0f9d58", "warn": "#b7791f", "bad": "#c53030"},
    "midnight": {"label": "Midnight Indigo", "bg1": "rgba(99,102,241,0.16)", "bg2": "rgba(168,85,247,0.12)", "base1": "#0a0f1e", "base2": "#11172a", "base3": "#151d35", "accent": "#c4b5fd", "accent2": "#7dd3fc", "ok": "#86efac", "warn": "#fde68a", "bad": "#fda4af"},
    "obsidian": {"label": "Obsidian Luxe", "bg1": "rgba(244,114,182,0.10)", "bg2": "rgba(250,204,21,0.08)", "base1": "#0b0b10", "base2": "#14141d", "base3": "#1c1d28", "accent": "#f9a8d4", "accent2": "#fcd34d", "ok": "#4ade80", "warn": "#facc15", "bad": "#fb7185"},
}

ML_MODEL_SPECS = [
    ("use_lr", "LR", "Logistic Regression"),
    ("use_rf", "RF", "Random Forest"),
    ("use_svm", "SVM", "Support Vector Machine"),
    ("use_xgb", "XGBoost", "Extreme Gradient Boosting"),
    ("use_dt", "DT", "Decision Tree"),
]
DL_MODEL_SPECS = [
    ("use_mlp", "MLP", "Shallow multilayer perceptron"),
    ("use_dnn", "DNN", "Deep dense neural network"),
    ("use_cnn1d", "CNN1D", "1D convolutional network"),
    ("use_lstm", "LSTM", "Sequence modelling with LSTM"),
    ("use_aae", "AAE*", "Autoencoder-regularized classifier (experimental)"),
    ("use_cnn_lstm", "CNN-LSTM", "Conv + LSTM hybrid"),
    ("use_cnn_bilstm", "CNN-BiLSTM", "Conv + bidirectional LSTM"),
    ("use_cnn_gru", "CNN-GRU", "Conv + GRU hybrid"),
    ("use_gnn", "GNN", "Graph neural network"),
    ("use_gat", "GAT", "Graph attention network"),
    ("use_transformer", "Transformer", "Tabular transformer encoder"),
    ("use_transfer_learning", "Transfer Learning*", "Self-supervised encoder transfer (experimental)"),
]

def coerce_text(value, fallback: str = "") -> str:
    """Return a safe string for logs/tracebacks even when DB values are NaN/float."""
    try:
        if value is None:
            return fallback
        if isinstance(value, str):
            s = value.strip()
            return s if s else fallback
        if pd.isna(value):
            return fallback
        if isinstance(value, (bytes, bytearray)):
            s = value.decode("utf-8", errors="ignore").strip()
            return s if s else fallback
        s = str(value).strip()
        return s if s and s.lower() != "nan" else fallback
    except Exception:
        return fallback

def tail_text(value, limit: int = 12000, fallback: str = "(tanpa log)") -> str:
    s = coerce_text(value, fallback=fallback)
    return s[-limit:] if len(s) > limit else s

@st.cache_data(show_spinner=False, ttl=60)
def get_tensorflow_health() -> Dict:
    info = {
        "ok": False,
        "import_ok": False,
        "version": None,
        "message": "",
        "gpu_devices": [],
        "cpu_devices": [],
        "traceback": "",
    }
    try:
        import tensorflow as tf
        info["import_ok"] = True
        info["version"] = getattr(tf, "__version__", None)
        try:
            gpus = [d.name for d in tf.config.list_physical_devices("GPU")]
            cpus = [d.name for d in tf.config.list_physical_devices("CPU")]
            info["gpu_devices"] = gpus
            info["cpu_devices"] = cpus
            info["ok"] = True
            gpu_msg = f"GPU: {len(gpus)}" if gpus else "GPU: 0"
            cpu_msg = f"CPU: {len(cpus)}" if cpus else "CPU: 0"
            info["message"] = f"TensorFlow {info['version']} siap. {cpu_msg} · {gpu_msg}"
        except Exception as exc:
            info["message"] = f"TensorFlow berhasil di-import, tetapi device query gagal: {type(exc).__name__}: {exc}"
            info["traceback"] = traceback.format_exc()
    except Exception as exc:
        info["message"] = f"{type(exc).__name__}: {exc}"
        info["traceback"] = traceback.format_exc()
    return info


def render_tf_health_panel(tf_health: Dict, compact: bool = False):
    st.markdown("### ⚙️ TensorFlow / DL runtime health")
    msg = coerce_text(tf_health.get("message"), fallback="Tidak ada status runtime.")
    if tf_health.get("ok"):
        st.success(msg)
        cols = st.columns(3)
        cols[0].metric("TensorFlow", tf_health.get("version") or "-")
        cols[1].metric("CPU devices", len(tf_health.get("cpu_devices") or []))
        cols[2].metric("GPU devices", len(tf_health.get("gpu_devices") or []))
    else:
        st.warning("Runtime DL belum siap. Dashboard tetap bisa dipakai dalam mode ML-only.")
        st.caption(msg)
        if not compact:
            tb = coerce_text(tf_health.get("traceback"))
            if tb:
                with st.expander("Lihat detail error TensorFlow", expanded=False):
                    st.code(tail_text(tb, limit=12000, fallback="(tanpa traceback)"), language="text")


def run_startup_self_test(pipeline_path: str, tf_health: Optional[Dict] = None) -> Dict:
    results: List[Tuple[str, str, str]] = []

    def add(status: str, title: str, detail: str):
        results.append((status, title, detail))

    p = Path(str(pipeline_path or ""))
    if p.exists():
        add("pass", "Pipeline path", f"Pipeline ditemukan: {p}")
        try:
            py_compile.compile(str(p), doraise=True)
            add("pass", "Pipeline syntax", "File pipeline lolos compile/syntax check.")
        except Exception as exc:
            add("fail", "Pipeline syntax", coerce_text(exc, fallback="Compile gagal."))
    else:
        add("fail", "Pipeline path", f"Pipeline tidak ditemukan: {p}")

    try:
        ARTIFACT_ROOT.mkdir(parents=True, exist_ok=True)
        JOB_ROOT.mkdir(parents=True, exist_ok=True)
        add("pass", "Artifact directories", "Folder artifact/job siap digunakan.")
    except Exception as exc:
        add("fail", "Artifact directories", coerce_text(exc, fallback="Folder artifact/job gagal dibuat."))

    try:
        init_db()
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("SELECT 1")
        add("pass", "Registry database", f"SQLite registry siap: {DB_PATH.name}")
    except Exception as exc:
        add("fail", "Registry database", coerce_text(exc, fallback="Registry DB tidak siap."))

    tf = tf_health or {}
    if tf.get("ok"):
        add("pass", "DL runtime", coerce_text(tf.get("message"), fallback="TensorFlow/DL runtime siap."))
    else:
        add("warn", "DL runtime", coerce_text(tf.get("message"), fallback="TensorFlow/DL runtime belum siap; fallback ML-only dapat dipakai."))

    n_fail = sum(1 for s, _, _ in results if s == 'fail')
    n_warn = sum(1 for s, _, _ in results if s == 'warn')
    overall = 'fail' if n_fail else ('warn' if n_warn else 'pass')
    return {'overall': overall, 'results': results, 'n_fail': n_fail, 'n_warn': n_warn}


def render_startup_self_test_panel(report: Dict, key_prefix: str = 'startup'):
    if not report:
        return
    overall = report.get('overall', 'warn')
    if overall == 'pass':
        st.success('Startup self-check: PASS')
    elif overall == 'warn':
        st.warning('Startup self-check: PASS with warnings')
    else:
        st.error('Startup self-check: FAIL')
    with st.expander('Startup self-check details', expanded=(overall != 'pass')):
        for i, (status, title, detail) in enumerate(report.get('results', [])):
            if status == 'pass':
                st.markdown(f"✅ **{title}** — {detail}")
            elif status == 'warn':
                st.markdown(f"⚠️ **{title}** — {detail}")
            else:
                st.markdown(f"❌ **{title}** — {detail}")

MODEL_COSTS = {
    "use_lr": 2, "use_dt": 2, "use_rf": 4, "use_svm": 5, "use_xgb": 5,
    "use_mlp": 4, "use_dnn": 5, "use_cnn1d": 6, "use_lstm": 7,
    "use_cnn_lstm": 8, "use_cnn_bilstm": 8, "use_cnn_gru": 8,
    "use_gnn": 9, "use_gat": 10, "use_transformer": 9,
    "use_aae": 9, "use_transfer_learning": 10,
}

ACADEMIC_GUARDS = [
    ("No test leakage", "Semua preprocessing penting dipelajari dari TRAIN, lalu diterapkan ke validation/test."),
    ("Validation for tuning", "Validation dipakai untuk early stopping, tuning, dan pemilihan model, bukan untuk pelaporan final."),
    ("Hold-out final test", "Test set dipertahankan sebagai evaluasi final yang tidak disentuh pada tahap fitting."),
    ("Train-only balancing", "Balancing/resampling hanya diterapkan pada TRAIN agar evaluasi tetap fair."),
    ("Transparent protocol", "Konfigurasi, split, scaler, balancing, copula, dan expert selection dicatat dalam manifest."),
    ("Reproducibility-first", "Run timestamp, parameter, output figure, log, dan summary dapat diunduh kembali untuk audit."),
]

DEMO_LOGIN_HINT = "Login lokal default: username `admin`, password `nids123`."
UPLOAD_HELP = {
    "nslkdd": {"train": "Upload KDDTrain+.txt", "test": "Upload KDDTest+.txt"},
    "unsw": {"train": "Upload UNSW_NB15_training-set.csv", "test": "Upload UNSW_NB15_testing-set.csv"},
    "insdn": {"normal": "Upload file normal traffic", "ovs": "Upload file OVS attack", "metasploit": "Upload file Metasploit"},
    "cicids": {"files": "Upload semua file CSV CICIDS2017. Bisa 8 file sekaligus."},
}

MODEL_FLAGS = [k for k, _, _ in (ML_MODEL_SPECS + DL_MODEL_SPECS + SEMI_ML_MODEL_SPECS + SEMI_DL_MODEL_SPECS + UNSUP_ML_MODEL_SPECS + UNSUP_DL_MODEL_SPECS + RL_ML_MODEL_SPECS + RL_DL_MODEL_SPECS)]

HYPERPARAM_DEFAULTS = {
    "lr_c": 1.0,
    "lr_solver": "lbfgs",
    "rf_n_estimators": 300,
    "rf_max_depth": 0,
    "svm_c": 1.0,
    "svm_kernel": "rbf",
    "xgb_n_estimators": 300,
    "xgb_max_depth": 6,
    "xgb_learning_rate": 0.05,
    "dt_max_depth": 18,
    "dl_epochs": 30,
    "dl_batch_size": 256,
}


@dataclass
class PreparedDatasetInput:
    mode: str
    paths: Dict[str, str]
    saved_files: List[Path]


class StreamlitLogWriter(io.StringIO):
    def __init__(self, placeholder, max_chars: int = 100000, refresh_sec: float = 0.12):
        super().__init__()
        self.placeholder = placeholder
        self.max_chars = max_chars
        self.refresh_sec = refresh_sec
        self._buffer = ""
        self._last_render = 0.0

    def write(self, s: str) -> int:
        if not s:
            return 0
        self._buffer += s
        if len(self._buffer) > self.max_chars:
            self._buffer = self._buffer[-self.max_chars:]
        now = time.time()
        if now - self._last_render >= self.refresh_sec or s.endswith("\n"):
            self.placeholder.code(self._buffer, language="text")
            self._last_render = now
        return len(s)

    def flush(self):
        self.placeholder.code(self._buffer, language="text")

    @property
    def value(self) -> str:
        return self._buffer


class NullPlaceholder:
    def code(self, *args, **kwargs):
        return None


class FileLogWriter(io.StringIO):
    def __init__(self, log_path: Path, max_chars: int = 200000):
        super().__init__()
        self.log_path = Path(log_path)
        self.log_path.parent.mkdir(parents=True, exist_ok=True)
        self._buffer = ""
        self.max_chars = max_chars

    def write(self, s: str) -> int:
        if not s:
            return 0
        with self.log_path.open("a", encoding="utf-8", errors="ignore") as f:
            f.write(s)
        self._buffer += s
        if len(self._buffer) > self.max_chars:
            self._buffer = self._buffer[-self.max_chars:]
        return len(s)

    def flush(self):
        return None

    @property
    def value(self) -> str:
        return self._buffer


def write_job_status(job_dir: Path, payload: Dict):
    job_dir.mkdir(parents=True, exist_ok=True)
    payload = dict(payload)
    payload.setdefault("updated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    (job_dir / "status.json").write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")


def read_job_status(job_dir: Path) -> Dict:
    path = Path(job_dir) / "status.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def is_pid_alive(pid: Optional[int]) -> bool:
    if not pid:
        return False
    try:
        os.kill(int(pid), 0)
    except OSError:
        return False
    except Exception:
        return False
    return True


def _db_insert_job(run_id: str, status: str, config: Dict, audit: Optional[Dict], guards: Dict, job_dir: Optional[Path] = None, job_pid: Optional[int] = None, error_text: str = ""):
    ml_active, dl_active = active_model_groups(config)
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            INSERT OR REPLACE INTO experiments (
                run_id, created_at, updated_at, status, experiment_name, dataset, classification_mode, workflow_mode, experiment_type,
                scaler, balance, copula, cv_folds, top_k, seed, complexity_score, complexity_label, notes, ml_models, dl_models,
                theme, protocol_locked, pipeline_path, config_json, audit_json, guards_json, error_text, job_pid, stop_requested, job_dir
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                run_id, created_at, created_at, status, config.get("experiment_name"), config.get("dataset"),
                "Binary" if config.get("binary", True) else "Multiclass", config.get("workflow_mode"), config.get("experiment_type"),
                config.get("resolved_scaler", config.get("scaler")), config.get("balance"), config.get("resolved_copula", config.get("copula_family")),
                int(config.get("cv_folds", 0)), int(config.get("top_k", 0)), int(config.get("seed", 0)),
                int(config.get("complexity_score", 0)), config.get("complexity_label"), config.get("notes", ""),
                ", ".join(ml_active), ", ".join(dl_active), THEME_OPTIONS.get(config.get("theme", "cyber"), THEME_OPTIONS["cyber"])["label"],
                1 if config.get("protocol_lock") else 0, config.get("pipeline_path"), json.dumps(config, default=str), json.dumps(audit or {}, default=str), json.dumps(guards, default=str),
                error_text, None if job_pid is None else int(job_pid), 0, str(job_dir) if job_dir else None
            )
        )
        conn.commit()


def db_insert_queued(run_id: str, config: Dict, audit: Optional[Dict], guards: Dict, job_dir: Path):
    _db_insert_job(run_id, "queued", config, audit, guards, job_dir=job_dir, job_pid=None)


def db_mark_running(run_id: str, pid: int):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("UPDATE experiments SET updated_at=?, status=?, job_pid=?, stop_requested=0 WHERE run_id=?", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "running", int(pid), run_id))
        conn.commit()


def db_request_stop(run_id: str):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("UPDATE experiments SET updated_at=?, stop_requested=1 WHERE run_id=?", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), run_id))
        conn.commit()


def enqueue_background_job(config: Dict, audit: Optional[Dict], guards: Dict) -> str:
    run_id = make_run_id(config)
    job_dir = JOB_ROOT / run_id
    job_dir.mkdir(parents=True, exist_ok=True)
    (job_dir / "config.json").write_text(json.dumps(config, indent=2, default=str), encoding="utf-8")
    (job_dir / "audit.json").write_text(json.dumps(audit or {}, indent=2, default=str), encoding="utf-8")
    (job_dir / "guards.json").write_text(json.dumps(guards or {}, indent=2, default=str), encoding="utf-8")
    write_job_status(job_dir, {"state": "queued", "run_id": run_id, "dataset": config.get("dataset")})
    db_insert_queued(run_id, config, audit, guards, job_dir)
    return run_id


def launch_worker_process(job_dir: Path):
    return subprocess.Popen([sys.executable, str(Path(__file__).resolve()), "--worker", str(job_dir)], cwd=str(APP_DIR), stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, start_new_session=True)


def reconcile_background_jobs():
    df = load_registry_df()
    if df.empty or "status" not in df.columns:
        return
    running = df[df["status"] == "running"].copy()
    for _, row in running.iterrows():
        pid = row.get("job_pid")
        if is_pid_alive(pid):
            continue
        job_dir = Path(row.get("job_dir")) if row.get("job_dir") else None
        status_info = read_job_status(job_dir) if job_dir else {}
        state = status_info.get("state")
        if state in {"success", "failed", "stopped"}:
            continue
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("UPDATE experiments SET updated_at=?, status=?, error_text=? WHERE run_id=?", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "failed", row.get("error_text") or "Worker terminated unexpectedly.", row.get("run_id")))
            conn.commit()
        if job_dir:
            write_job_status(job_dir, {"state": "failed", "run_id": row.get("run_id"), "error": "Worker terminated unexpectedly."})


def start_next_queued_job() -> Optional[str]:
    df = load_registry_df()
    if df.empty:
        return None
    running = df[df["status"] == "running"].copy()
    for _, row in running.iterrows():
        if is_pid_alive(row.get("job_pid")):
            return None
    queued = df[df["status"] == "queued"].copy()
    if queued.empty:
        return None
    queued = queued.sort_values(["created_at", "run_id"], ascending=[True, True])
    row = queued.iloc[0]
    job_dir = Path(row["job_dir"])
    proc = launch_worker_process(job_dir)
    db_mark_running(str(row["run_id"]), proc.pid)
    write_job_status(job_dir, {"state": "running", "run_id": str(row["run_id"]), "pid": proc.pid})
    return str(row["run_id"])


def stop_background_job(run_id: str) -> bool:
    df = load_registry_df()
    if df.empty:
        return False
    rows = df[df["run_id"] == run_id]
    if rows.empty:
        return False
    row = rows.iloc[0]
    db_request_stop(run_id)
    pid = row.get("job_pid")
    job_dir = Path(row.get("job_dir")) if row.get("job_dir") else None
    try:
        if pid and is_pid_alive(pid):
            try:
                os.killpg(int(pid), signal.SIGTERM)
            except Exception:
                os.kill(int(pid), signal.SIGTERM)
    except Exception:
        pass
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("UPDATE experiments SET updated_at=?, status=?, error_text=? WHERE run_id=?", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "stopped", "Stopped by user request.", run_id))
        conn.commit()
    if job_dir:
        write_job_status(job_dir, {"state": "stopped", "run_id": run_id, "error": "Stopped by user request."})
    return True


def hydrate_result_bundle_from_row(row_like) -> Optional[Dict]:
    row = dict(row_like) if not isinstance(row_like, dict) else row_like
    artifact_root = Path(row.get("output_root")) if row.get("output_root") else None
    if not artifact_root or not artifact_root.exists():
        return None
    summary_path = artifact_root / "summary.csv"
    summary_df = pd.read_csv(summary_path) if summary_path.exists() else pd.DataFrame()
    logs_path = artifact_root / "live_logs.txt"
    logs = logs_path.read_text(encoding="utf-8", errors="ignore") if logs_path.exists() else coerce_text(row.get("error_text"), "")
    fig_paths = sorted((artifact_root / "figures").glob("*.png")) if (artifact_root / "figures").exists() else []
    log_paths = sorted((artifact_root / "logs").glob("*.log")) if (artifact_root / "logs").exists() else []
    config = json.loads(row.get("config_json") or "{}")
    preprep_path = artifact_root / "prepreprocess_report.json"
    preprep_report = None
    if preprep_path.exists():
        try:
            preprep_report = json.loads(preprep_path.read_text(encoding="utf-8"))
        except Exception:
            preprep_report = None
    return {
        "run_id": row.get("run_id"),
        "timestamp": row.get("updated_at") or row.get("created_at"),
        "result": {"summary": summary_df, "prepreprocess_report": preprep_report},
        "logs": logs,
        "elapsed_sec": None,
        "fig_paths": fig_paths,
        "log_paths": log_paths,
        "config": config,
        "pipeline_path": config.get("pipeline_path"),
        "artifact_zip_path": row.get("artifact_zip_path"),
        "artifact_root": row.get("output_root"),
    }


def load_latest_success_bundle() -> Optional[Dict]:
    df = load_registry_df()
    if df.empty or "status" not in df.columns:
        return None
    success = df[df["status"] == "success"].copy()
    if success.empty:
        return None
    row = success.sort_values(["updated_at", "created_at"], ascending=[False, False]).iloc[0]
    return hydrate_result_bundle_from_row(row)


def worker_main(job_dir_arg: str):
    job_dir = Path(job_dir_arg)
    config = json.loads((job_dir / "config.json").read_text(encoding="utf-8"))
    audit = json.loads((job_dir / "audit.json").read_text(encoding="utf-8")) if (job_dir / "audit.json").exists() else {}
    guards = json.loads((job_dir / "guards.json").read_text(encoding="utf-8")) if (job_dir / "guards.json").exists() else {}
    run_id = job_dir.name
    db_mark_running(run_id, os.getpid())
    write_job_status(job_dir, {"state": "running", "run_id": run_id, "pid": os.getpid()})
    log_path = job_dir / "worker_live.log"
    file_writer = FileLogWriter(log_path)
    try:
        module, module_path = load_pipeline_module(config["pipeline_path"])
        start = time.time()
        with contextlib.redirect_stdout(file_writer), contextlib.redirect_stderr(file_writer):
            pipeline_cls = getattr(module, "NIDSPipeline")
            run_args = dict(
                dataset=config["dataset"], paths=config["paths"], binary=config["binary"], copula_family=config["copula_family"],
                balance=config["balance"], scaler=config["scaler"], top_k=config["top_k"], cv_folds=config["cv_folds"],
                validation_method=config.get("validation_method", "holdout"), validation_repeats=config.get("validation_repeats", 3), bootstrap_rounds=config.get("bootstrap_rounds", 30),
                nested_inner_folds=config.get("nested_inner_folds", 3), feature_method=config.get("feature_method", "mutual_info"), optimize_ml=config.get("optimize_ml", False),
                optimize_dl=config.get("optimize_dl", False), ensemble_method=config.get("ensemble_method", "weighted_soft_vote"), hybrid_method=config.get("hybrid_method", "weighted_ml_dl"),
                use_ensemble=config.get("use_ensemble", True), use_hybrid=config.get("use_hybrid", True),
                use_lr=config.get("use_lr", True), use_rf=config.get("use_rf", True), use_svm=config.get("use_svm", True), use_xgb=config.get("use_xgb", True), use_dt=config.get("use_dt", True),
                use_mlp=config.get("use_mlp", True), use_dnn=config.get("use_dnn", False), use_cnn1d=config.get("use_cnn1d", False), use_lstm=config.get("use_lstm", False),
                use_aae=config.get("use_aae", False), use_cnn_lstm=config.get("use_cnn_lstm", False), use_cnn_bilstm=config.get("use_cnn_bilstm", True), use_cnn_gru=config.get("use_cnn_gru", True),
                use_gnn=config.get("use_gnn", False), use_gat=config.get("use_gat", False), use_transformer=config.get("use_transformer", True), use_transfer_learning=config.get("use_transfer_learning", False),
                use_bn=config["use_bn"], use_shap=config["use_shap"], sample_frac=config["sample_frac"], cicids_split=config["cicids_split"],
                model_params=config.get("model_params", {}), training_params=config.get("training_params", {}),
                learning_paradigm=config.get("learning_paradigm", "supervised"),
                use_label_spreading=config.get("use_label_spreading", False), use_self_training_svm=config.get("use_self_training_svm", False),
                use_self_training_rf=config.get("use_self_training_rf", False), use_pseudo_label_dnn=config.get("use_pseudo_label_dnn", False),
                use_isolation_forest=config.get("use_isolation_forest", False), use_oneclass_svm_anom=config.get("use_oneclass_svm_anom", False),
                use_lof=config.get("use_lof", False), use_pca_anom=config.get("use_pca_anom", False), use_kmeans_anom=config.get("use_kmeans_anom", False),
                use_autoencoder_anom=config.get("use_autoencoder_anom", False), use_rl_threshold=config.get("use_rl_threshold", False),
                semi_label_fraction=config.get("semi_label_fraction", 0.3), pseudo_label_threshold=config.get("pseudo_label_threshold", 0.9),
                anomaly_contamination=config.get("anomaly_contamination", 0.1), rl_episodes=config.get("rl_episodes", 40),
            )
            result = pipeline_cls(**run_args).run()
        elapsed = time.time() - start
        fig_paths, log_paths = collect_output_files(module_path.parent, config["dataset"])
        latest = {
            "run_id": run_id,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "result": result,
            "logs": log_path.read_text(encoding="utf-8", errors="ignore") if log_path.exists() else "",
            "elapsed_sec": elapsed,
            "fig_paths": fig_paths,
            "log_paths": log_paths,
            "config": config,
            "pipeline_path": str(module_path),
        }
        manifest = make_methodology_manifest(latest, audit=audit, guards=guards)
        report_md = make_academic_report_markdown(latest, config=config, audit=audit, guards=guards, style=config.get("report_template", "formal_dissertation"))
        out_dir, zip_path = save_artifacts(run_id, latest, audit, guards, manifest, report_md)
        latest["artifact_zip_path"] = str(zip_path)
        latest["artifact_root"] = str(out_dir)
        db_complete_run(run_id, status="success", config=config, manifest=manifest, audit=audit, guards=guards, result_bundle=latest, artifact_zip_path=zip_path, output_root=out_dir)
        best_model, best_f1, best_acc, best_auc = get_best_model_info(result.get("summary") if isinstance(result, dict) else None)
        write_job_status(job_dir, {"state": "success", "run_id": run_id, "artifact_zip_path": str(zip_path), "output_root": str(out_dir), "best_model": best_model, "f1": best_f1, "accuracy": best_acc, "auc": best_auc})
    except Exception as exc:
        err = traceback.format_exc()
        (job_dir / "worker_error.txt").write_text(err, encoding="utf-8")
        db_complete_run(run_id, status="failed", config=config, manifest=None, audit=audit, guards=guards, result_bundle=None, artifact_zip_path=None, output_root=None, error_text=err)
        write_job_status(job_dir, {"state": "failed", "run_id": run_id, "error": str(exc)})
        return 1
    return 0


def init_db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    JOB_ROOT.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS experiments (
                run_id TEXT PRIMARY KEY,
                created_at TEXT,
                updated_at TEXT,
                status TEXT,
                experiment_name TEXT,
                dataset TEXT,
                classification_mode TEXT,
                workflow_mode TEXT,
                experiment_type TEXT,
                scaler TEXT,
                balance TEXT,
                copula TEXT,
                cv_folds INTEGER,
                top_k INTEGER,
                seed INTEGER,
                complexity_score INTEGER,
                complexity_label TEXT,
                best_model TEXT,
                f1 REAL,
                accuracy REAL,
                auc REAL,
                notes TEXT,
                ml_models TEXT,
                dl_models TEXT,
                theme TEXT,
                protocol_locked INTEGER,
                pipeline_path TEXT,
                config_json TEXT,
                manifest_json TEXT,
                audit_json TEXT,
                guards_json TEXT,
                artifact_zip_path TEXT,
                output_root TEXT,
                error_text TEXT,
                job_pid INTEGER,
                stop_requested INTEGER DEFAULT 0,
                job_dir TEXT
            )
            """
        )
        existing = {row[1] for row in conn.execute("PRAGMA table_info(experiments)").fetchall()}
        for col_def in [
            ("job_pid", "INTEGER"),
            ("stop_requested", "INTEGER DEFAULT 0"),
            ("job_dir", "TEXT"),
        ]:
            if col_def[0] not in existing:
                conn.execute(f"ALTER TABLE experiments ADD COLUMN {col_def[0]} {col_def[1]}")
        conn.commit()

def init_state():
    init_db()
    st.session_state.setdefault("result_bundle", None)
    st.session_state.setdefault("history", [])
    st.session_state.setdefault("run_counter", 0)
    st.session_state.setdefault("authenticated", False)
    st.session_state.setdefault("auth_user", "")
    st.session_state.setdefault("theme_choice", "cyber")
    st.session_state.setdefault("selected_compare_runs", [])
    st.session_state.setdefault("selected_registry_run", None)
    st.session_state.setdefault("tf_health_snapshot", None)


def get_login_credentials() -> Tuple[str, str]:
    return os.getenv("NIDS_GUI_USER", "admin"), os.getenv("NIDS_GUI_PASS", "nids123")


def render_login_gate():
    if st.session_state.get("authenticated"):
        return
    expected_user, expected_pass = get_login_credentials()
    st.markdown(
        """
        <div class="login-shell">
            <h1>🔐 Hybrid NIDS Research Studio v4.6.1 Final Study</h1>
            <p>Masuk terlebih dahulu untuk membuka dashboard eksperimen. Versi ini menekankan audit trail, protocol lock, dan registry eksperimen yang persisten.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    left, center, right = st.columns([1.1, 1.5, 1.1])
    with center:
        with st.form("login_form", clear_on_submit=False):
            username = st.text_input("Username", value=st.session_state.get("auth_user", ""))
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Login", use_container_width=True, type="primary")
        st.caption(DEMO_LOGIN_HINT)
        if submitted:
            user_ok = hmac.compare_digest(username.strip(), expected_user)
            pass_ok = hmac.compare_digest(password, expected_pass)
            if user_ok and pass_ok:
                st.session_state["authenticated"] = True
                st.session_state["auth_user"] = username.strip()
                st.success(f"Selamat datang, {username.strip()}.")
                st.rerun()
            else:
                st.error("Username atau password tidak cocok.")
    st.stop()


def render_account_box():
    user = st.session_state.get("auth_user") or "guest"
    theme_key = st.session_state.get("theme_choice", "cyber")
    theme_label = THEME_OPTIONS.get(theme_key, THEME_OPTIONS["cyber"])["label"]
    st.sidebar.markdown("---")
    st.sidebar.markdown(f"**User:** {user}")
    st.sidebar.caption(f"Theme aktif: {theme_label}")
    st.sidebar.caption(f"Registry DB: `{DB_PATH.name}`")
    if st.sidebar.button("Logout", use_container_width=True, key="logout_btn_v2"):
        st.session_state["authenticated"] = False
        st.session_state["auth_user"] = ""
        st.rerun()


def inject_css(theme_key: str = "cyber"):
    theme = THEME_OPTIONS.get(theme_key, THEME_OPTIONS["cyber"])
    is_light = theme_key == "academic"
    text_color = "#10233f" if is_light else "#eef6ff"
    muted_color = "#38506f" if is_light else "#b7c8df"
    card_bg = "rgba(255,255,255,0.78)" if is_light else "rgba(255,255,255,0.045)"
    border = "rgba(21,43,77,0.08)" if is_light else "rgba(255,255,255,0.08)"
    code_bg = "#f7f9fc" if is_light else "#0d1528"
    sidebar_bg = "rgba(255,255,255,0.86)" if is_light else "rgba(13,21,40,0.92)"
    hero_overlay = "rgba(29,78,216,0.12)" if is_light else "rgba(68,114,255,0.26)"
    hero_overlay2 = "rgba(91,124,250,0.10)" if is_light else "rgba(25,198,255,0.16)"
    style = f"""
        <style>
        .stApp {{
            color: {text_color};
            background:
                radial-gradient(circle at top left, {theme['bg1']}, transparent 30%),
                radial-gradient(circle at top right, {theme['bg2']}, transparent 26%),
                linear-gradient(180deg, {theme['base1']} 0%, {theme['base2']} 42%, {theme['base3']} 100%);
        }}
        .main > div {{padding-top: 1.2rem;}}
        .glass {{
            background: {card_bg};
            border: 1px solid {border};
            box-shadow: 0 18px 50px rgba(0,0,0,0.25);
            backdrop-filter: blur(10px);
            -webkit-backdrop-filter: blur(10px);
            border-radius: 24px;
            padding: 18px 20px;
        }}
        .hero {{
            background: linear-gradient(135deg, {hero_overlay}, {hero_overlay2}), {card_bg};
            border: 1px solid color-mix(in srgb, {theme['accent2']} 22%, transparent);
            border-radius: 28px;
            padding: 26px 28px;
            margin-bottom: 1rem;
            box-shadow: 0 22px 60px rgba(0,0,0,0.22);
        }}
        .hero h1 {{font-size: 2.2rem; margin: 0 0 0.35rem 0; color: {text_color};}}
        .hero p {{opacity: 0.95; line-height: 1.55; margin-bottom: 0.7rem; color: {muted_color};}}
        .pill {{display:inline-block; padding:0.24rem 0.72rem; margin:0.18rem 0.35rem 0.15rem 0; border-radius:999px; font-size:0.82rem; background: color-mix(in srgb, {theme['accent2']} 16%, transparent); border:1px solid color-mix(in srgb, {theme['accent2']} 22%, transparent);}}
        .metric-card {{background: {card_bg}; border:1px solid {border}; border-radius:22px; padding:15px 16px; min-height:115px; box-shadow: inset 0 1px 0 rgba(255,255,255,0.05);}}
        .metric-label {{font-size:0.86rem; opacity:0.78; margin-bottom:0.5rem;}}
        .metric-value {{font-size:1.42rem; font-weight:800; margin-bottom:0.35rem;}}
        .metric-note {{font-size:0.82rem; opacity:0.72; line-height:1.35;}}
        .mini-card {{background:{card_bg}; border-radius:18px; border:1px solid {border}; padding:12px 14px; min-height:95px;}}
        .small-muted {{font-size:0.82rem; opacity:0.82; color:{muted_color};}}
        .leader-row {{background: color-mix(in srgb, {card_bg} 90%, transparent); border:1px solid {border}; padding:10px 12px; border-radius:16px; margin-bottom:0.45rem;}}
        .audit-ok {{color:{theme['ok']}; font-weight:700;}} .audit-warn {{color:{theme['warn']}; font-weight:700;}} .audit-bad {{color:{theme['bad']}; font-weight:700;}}
        .readiness-good {{color:{theme['ok']}; font-weight:700;}} .readiness-mid {{color:{theme['warn']}; font-weight:700;}} .readiness-bad {{color:{theme['bad']}; font-weight:700;}}
        .stCode, code, pre {{background:{code_bg} !important;}}
        .stSidebar > div:first-child {{background:{sidebar_bg};}}
        .stTabs [data-baseweb="tab"] {{font-weight:700;}}
        .stButton > button, .stDownloadButton > button, [data-testid="stFormSubmitButton"] button {{
            border-radius:14px !important;
            border:1px solid color-mix(in srgb, {theme['accent2']} 22%, transparent) !important;
            background: linear-gradient(135deg, color-mix(in srgb, {theme['accent']} 14%, transparent), color-mix(in srgb, {theme['accent2']} 18%, transparent)) !important;
            color:{text_color} !important; font-weight:700 !important; box-shadow:0 10px 24px rgba(0,0,0,0.12);
        }}
        .login-shell {{max-width:620px; margin:4vh auto 0 auto; padding:28px; background:{card_bg}; border:1px solid {border}; border-radius:28px; box-shadow:0 24px 70px rgba(0,0,0,0.20);}}
        .footer-note {{margin-top:1rem; opacity:0.78; font-size:0.84rem; color:{muted_color};}}
        </style>
    """
    st.markdown(style, unsafe_allow_html=True)


def metric_card(label: str, value: str, note: str = ""):
    st.markdown(
        f"<div class='metric-card'><div class='metric-label'>{label}</div><div class='metric-value'>{value}</div><div class='metric-note'>{note}</div></div>",
        unsafe_allow_html=True,
    )


def mini_card(label: str, value: str, note: str = ""):
    st.markdown(
        f"<div class='mini-card'><div class='metric-label'>{label}</div><div class='metric-value' style='font-size:1rem'>{value}</div><div class='metric-note'>{note}</div></div>",
        unsafe_allow_html=True,
    )


def resolve_default_pipeline_path() -> Path:
    if DEFAULT_PIPELINE_PATH.exists():
        return DEFAULT_PIPELINE_PATH
    return FALLBACK_PIPELINE_PATH


def load_pipeline_module(path: str):
    path_obj = Path(path)
    if not path_obj.exists():
        raise FileNotFoundError(f"Pipeline file tidak ditemukan: {path_obj}")
    module_name = f"nids_pipeline_module_{hashlib.md5(str(path_obj).encode()).hexdigest()[:8]}"
    spec = importlib.util.spec_from_file_location(module_name, path_obj)
    if spec is None or spec.loader is None:
        raise ImportError(f"Tidak dapat memuat module dari {path_obj}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module, path_obj


def path_status(path_str: str) -> Tuple[str, str]:
    if not path_str:
        return "⚪", "Belum diisi"
    p = Path(path_str)
    if p.exists():
        return "✅", "Ditemukan"
    return "❌", "Tidak ditemukan"


def render_path_status(paths: Dict[str, str], dataset: str):
    with st.sidebar.expander("Status path / upload", expanded=False):
        for key, value in paths.items():
            icon, label = path_status(value)
            st.markdown(f"**{key}** — {icon} {label}")
            if value:
                st.caption(value)
        st.caption(DATASET_NOTES.get(dataset, ""))


def save_uploaded_file(dataset: str, role: str, uploaded_file) -> Path:
    ts = int(time.time())
    target_dir = UPLOAD_ROOT / dataset / f"session_{ts}"
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / uploaded_file.name
    target.write_bytes(uploaded_file.getbuffer())
    return target


def save_multiple_uploaded_files(dataset: str, uploaded_files: List) -> List[Path]:
    ts = int(time.time())
    target_dir = UPLOAD_ROOT / dataset / f"session_{ts}"
    target_dir.mkdir(parents=True, exist_ok=True)
    saved = []
    for uf in uploaded_files:
        target = target_dir / uf.name
        target.write_bytes(uf.getbuffer())
        saved.append(target)
    return saved


def render_upload_inputs(dataset: str) -> PreparedDatasetInput:
    mode = st.sidebar.radio("Sumber data", options=["Path manual", "Upload file"], horizontal=False)
    saved_files: List[Path] = []
    paths: Dict[str, str] = {}
    if mode == "Path manual":
        for key, default_val in DEFAULT_PATHS[dataset].items():
            paths[key] = st.sidebar.text_input(f"{key}", value=default_val, key=f"path_{dataset}_{key}_v2")
        return PreparedDatasetInput(mode=mode, paths=paths, saved_files=saved_files)
    st.sidebar.caption("File yang di-upload akan disimpan sementara ke folder `uploaded_datasets/` di samping app Streamlit.")
    if dataset in {"nslkdd", "unsw", "insdn"}:
        for key in DEFAULT_PATHS[dataset].keys():
            uploaded = st.sidebar.file_uploader(f"Upload {key}", type=["csv", "txt"], key=f"upload_{dataset}_{key}_v2", help=UPLOAD_HELP[dataset].get(key, "Upload file dataset"))
            if uploaded is not None:
                saved = save_uploaded_file(dataset, key, uploaded)
                paths[key] = str(saved)
                saved_files.append(saved)
            else:
                paths[key] = ""
    else:
        uploaded_files = st.sidebar.file_uploader("Upload file CSV CICIDS2017", type=["csv"], key="upload_cicids_files_v2", accept_multiple_files=True, help=UPLOAD_HELP[dataset]["files"])
        if uploaded_files:
            saved = save_multiple_uploaded_files(dataset, uploaded_files)
            saved_files.extend(saved)
            paths["dir"] = str(saved[0].parent)
        else:
            paths["dir"] = ""
    return PreparedDatasetInput(mode=mode, paths=paths, saved_files=saved_files)


def validate_input_paths(dataset: str, paths: Dict[str, str]) -> Optional[str]:
    if dataset == "cicids":
        dir_path = Path(paths.get("dir", ""))
        if not str(dir_path):
            return "Path folder CICIDS2017 belum diisi."
        if not dir_path.exists():
            return f"Folder CICIDS2017 tidak ditemukan: {dir_path}"
        csvs = list(dir_path.glob("*.csv")) + list(dir_path.glob("*/*.csv"))
        if not csvs:
            return f"Tidak ada file CSV di folder CICIDS2017: {dir_path}"
        return None
    for key, value in paths.items():
        if not value:
            return f"Path `{key}` belum diisi."
        if not Path(value).exists():
            return f"Path `{key}` tidak ditemukan: {value}"
    return None


def count_data_rows(path: Path, has_header: bool) -> int:
    try:
        with path.open("r", encoding="utf-8", errors="ignore") as f:
            count = sum(1 for _ in f)
        if has_header and count > 0:
            count -= 1
        return max(count, 0)
    except Exception:
        return -1


def _read_sample(path: Path, dataset: str, nrows: int = 2000) -> pd.DataFrame:
    if dataset == "nslkdd":
        return pd.read_csv(path, header=None, nrows=nrows)
    return pd.read_csv(path, nrows=nrows, low_memory=False)


def _infer_label_counts(df: pd.DataFrame, dataset: str) -> Dict[str, int]:
    try:
        if dataset == "nslkdd":
            label_col = df.columns[-2] if len(df.columns) >= 43 else df.columns[-1]
            return df[label_col].astype(str).value_counts().head(8).to_dict()
        lowered = {str(c).strip().lower(): c for c in df.columns}
        if "attack_cat" in lowered:
            return df[lowered["attack_cat"]].astype(str).value_counts().head(8).to_dict()
        if "label" in lowered:
            return df[lowered["label"]].astype(str).value_counts().head(8).to_dict()
    except Exception:
        pass
    return {}


@st.cache_data(show_spinner=False)
def audit_dataset_cached(dataset: str, paths_json: str) -> Dict:
    paths = json.loads(paths_json)
    files: List[Tuple[str, Path]] = []
    if dataset == "cicids":
        dir_path = Path(paths.get("dir", ""))
        csvs = sorted(list(dir_path.glob("*.csv")) + list(dir_path.glob("*/*.csv"))) if dir_path.exists() else []
        files.extend([(p.name, p) for p in csvs])
    else:
        files.extend([(role, Path(path)) for role, path in paths.items() if path])

    file_reports = []
    total_rows = 0
    sample_frames = []
    for role, path in files:
        exists = path.exists()
        report = {
            "role": role,
            "path": str(path),
            "exists": exists,
            "rows_estimate": None,
            "sample_rows": 0,
            "sample_columns": 0,
            "sample_missing_pct": None,
            "sample_duplicate_pct": None,
            "label_preview": {},
            "error": None,
        }
        if not exists:
            report["error"] = "File not found"
            file_reports.append(report)
            continue
        has_header = dataset != "nslkdd"
        rows_est = count_data_rows(path, has_header=has_header)
        report["rows_estimate"] = rows_est
        total_rows += max(rows_est, 0)
        try:
            sample = _read_sample(path, dataset, nrows=1500)
            report["sample_rows"] = int(len(sample))
            report["sample_columns"] = int(sample.shape[1])
            report["sample_missing_pct"] = round(float(sample.isna().mean().mean() * 100), 4)
            report["sample_duplicate_pct"] = round(float(sample.duplicated().mean() * 100), 4)
            report["label_preview"] = _infer_label_counts(sample, dataset)
            sample_frames.append(sample.head(400))
        except Exception as exc:
            report["error"] = str(exc)
        file_reports.append(report)

    combined_sample = pd.concat(sample_frames, ignore_index=True) if sample_frames else pd.DataFrame()
    overall = {
        "dataset": dataset,
        "dataset_label": DATASET_LABELS.get(dataset, dataset),
        "n_files": len(files),
        "total_rows_estimate": total_rows,
        "sample_rows_combined": int(len(combined_sample)),
        "sample_columns_combined": int(combined_sample.shape[1]) if not combined_sample.empty else 0,
        "sample_missing_pct": round(float(combined_sample.isna().mean().mean() * 100), 4) if not combined_sample.empty else None,
        "sample_duplicate_pct": round(float(combined_sample.duplicated().mean() * 100), 4) if not combined_sample.empty else None,
    }
    return {"overall": overall, "files": file_reports}


def compute_guardrails(config: Dict, audit: Optional[Dict], validation_error: Optional[str]) -> Dict:
    issues = []
    score = 100

    def add(status: str, title: str, detail: str):
        nonlocal score
        issues.append({"status": status, "title": title, "detail": detail})
        if status == "warn":
            score -= 8
        elif status == "fail":
            score -= 18

    if validation_error:
        add("fail", "Input paths", validation_error)
    else:
        add("pass", "Input paths", "Semua path utama terverifikasi ada.")

    add("pass", "Train-only preprocessing", "Pipeline dikonfigurasi untuk fit encoder, imputer, scaler, feature selection, dan balancing hanya pada TRAIN.")
    add("pass", "Hold-out final test", SPLIT_PROTOCOLS.get(config["dataset"], "Split protocol dataset-specific."))

    if config.get("workflow_mode") == "publication" and config.get("cv_folds", 2) < 5:
        add("warn", "Publication CV depth", "Mode publication idealnya menggunakan CV >= 5 untuk hasil yang lebih kuat.")
    else:
        add("pass", "CV depth", f"CV folds = {config.get('cv_folds')}")

    if config["dataset"] == "cicids" and config.get("workflow_mode") == "publication" and config.get("cicids_split") != "temporal":
        add("warn", "Temporal split recommendation", "Untuk publikasi CICIDS2017, split temporal sering lebih kuat untuk mengurangi temporal leakage.")
    else:
        add("pass", "Split recommendation", "Konfigurasi split tidak menimbulkan red flag utama untuk workflow yang dipilih.")

    ml_active, dl_active = active_model_groups(config)
    if config.get("experiment_type") in {"hybrid_fusion", "copula_bn"} and (not ml_active or not dl_active):
        add("fail", "Hybrid claim support", "Eksperimen hybrid membutuhkan minimal satu model ML dan satu model DL aktif.")
    elif config.get("use_hybrid") and (not ml_active or not dl_active):
        add("warn", "Hybrid runtime guard", "Hybrid aktif tetapi hanya satu grup model yang tersedia. Runtime akan menonaktifkan hybrid secara otomatis agar run tetap stabil.")
    else:
        add("pass", "Hybrid evidence", "Komposisi model aktif sesuai narasi eksperimen yang dipilih.")

    if config.get("use_aae") or config.get("use_transfer_learning"):
        add("warn", "Experimental models", "AAE / Transfer Learning di studio ini bersifat experimental scaffold. Jelaskan dengan hati-hati bila dipakai untuk klaim utama.")
    else:
        add("pass", "Experimental models", "Tidak ada scaffold eksperimental aktif untuk run ini.")

    if not config.get("protocol_lock"):
        add("fail", "Protocol lock", "Checklist protocol lock belum dikonfirmasi.")
    else:
        add("pass", "Protocol lock", "Pengguna telah mengunci rancangan eksperimen sebelum menjalankan pipeline.")

    if audit and audit.get("overall", {}).get("sample_duplicate_pct") is not None and audit["overall"]["sample_duplicate_pct"] > 10:
        add("warn", "Duplicate ratio", f"Sample duplicate rate = {audit['overall']['sample_duplicate_pct']:.2f}% — pertimbangkan inspeksi deduplikasi lebih lanjut.")
    else:
        add("pass", "Duplicate ratio", "Tidak ada indikasi duplicate rate yang mencolok pada sample audit.")

    if audit and audit.get("overall", {}).get("sample_missing_pct") is not None and audit["overall"]["sample_missing_pct"] > 5:
        add("warn", "Missing data", f"Missing cell ratio sample = {audit['overall']['sample_missing_pct']:.2f}% — evaluasi kualitas data sebelum klaim final.")
    else:
        add("pass", "Missing data", "Missing ratio sample tidak menunjukkan masalah besar untuk baseline umum.")

    score = max(min(score, 100), 0)
    label = "Excellent" if score >= 88 else "Good" if score >= 72 else "Watchlist" if score >= 55 else "Critical"
    return {"score": score, "label": label, "checks": issues}


def compute_readiness(config: Dict, validation_error: Optional[str], audit: Optional[Dict], guards: Dict) -> Tuple[int, str, List[str]]:
    score = 100
    notes: List[str] = []
    if validation_error:
        score -= 35
        notes.append(validation_error)
    if config.get("experiment_type") == "custom":
        score -= 4
        notes.append("Custom design memerlukan validasi manual lebih teliti.")
    if config.get("workflow_mode") == "publication":
        notes.append("Mode publication aktif: audit trail dan manifest akan lebih lengkap.")
    if audit and audit.get("overall", {}).get("n_files", 0) == 0:
        score -= 10
        notes.append("Belum ada file terdeteksi untuk diaudit.")
    score = min(score, guards.get("score", 100))
    if score >= 88:
        label = "Ready"
    elif score >= 70:
        label = "Mostly ready"
    elif score >= 55:
        label = "Needs review"
    else:
        label = "Blocked"
    return score, label, notes


def active_model_groups(config: Dict) -> Tuple[List[str], List[str]]:
    ml_specs = ML_MODEL_SPECS + SEMI_ML_MODEL_SPECS + UNSUP_ML_MODEL_SPECS + RL_ML_MODEL_SPECS
    dl_specs = DL_MODEL_SPECS + SEMI_DL_MODEL_SPECS + UNSUP_DL_MODEL_SPECS + RL_DL_MODEL_SPECS
    ml = [label for key, label, _ in ml_specs if config.get(key)]
    dl = [label for key, label, _ in dl_specs if config.get(key)]
    return ml, dl


def get_template_overrides(experiment_type: str) -> Dict[str, bool]:
    all_false = {k: False for k in MODEL_FLAGS}
    if experiment_type == "single_baseline":
        all_false.update(use_rf=True, use_mlp=True)
    elif experiment_type == "ml_ensemble":
        for k, _, _ in ML_MODEL_SPECS:
            all_false[k] = True
    elif experiment_type == "dl_ensemble":
        for k, _, _ in DL_MODEL_SPECS:
            all_false[k] = k not in {"use_aae", "use_transfer_learning"}
    elif experiment_type == "hybrid_fusion":
        all_false.update(use_rf=True, use_xgb=True, use_svm=True, use_mlp=True, use_dnn=True, use_cnn1d=True, use_lstm=True, use_cnn_bilstm=True, use_transformer=True)
    elif experiment_type == "copula_bn":
        all_false.update(use_rf=True, use_xgb=True, use_svm=True, use_mlp=True, use_dnn=True, use_cnn1d=True, use_lstm=True, use_cnn_bilstm=True, use_cnn_gru=True, use_transformer=True)
    else:
        return {}
    return all_false


def apply_template_to_state(template: Dict[str, bool]):
    for key, value in template.items():
        st.session_state[f"model_{key}"] = value


def estimate_complexity(config: Dict) -> Tuple[int, str, str]:
    score = 12
    for key in MODEL_FLAGS:
        if config.get(key):
            score += MODEL_COSTS.get(key, 3)
    score += max(0, int(config.get("cv_folds", 2)) - 1) * 5
    score += max(0, int(config.get("top_k", 40) / 10) - 2)
    if config.get("dataset") == "cicids":
        score += int(float(config.get("sample_frac", 1.0)) * 20)
        if config.get("cicids_split") == "temporal":
            score += 5
    if config.get("use_shap"):
        score += 8
    if config.get("use_bn"):
        score += 4
    if config.get("workflow_mode") == "publication":
        score += 10
    if score <= 40:
        return score, "Light", "Cocok untuk iterasi cepat atau baseline eksploratif."
    if score <= 75:
        return score, "Moderate", "Masih nyaman untuk eksperimen penelitian normal."
    if score <= 115:
        return score, "Heavy", "Perlu perhatian pada runtime, terutama untuk dataset besar."
    return score, "Extreme", "Berpotensi sangat berat. Pertimbangkan pengurangan model, CV, atau sample fraction."


def get_best_model_info(summary_df: Optional[pd.DataFrame]) -> Tuple[str, Optional[float], Optional[float], Optional[float]]:
    if not isinstance(summary_df, pd.DataFrame) or summary_df.empty:
        return "-", None, None, None
    tmp = summary_df.copy()
    for c in ["F1", "Accuracy", "AUC-ROC"]:
        if c in tmp.columns:
            tmp[c] = pd.to_numeric(tmp[c], errors="coerce")
    metric = "F1" if "F1" in tmp.columns else ("Accuracy" if "Accuracy" in tmp.columns else None)
    if metric is None or tmp[metric].dropna().empty:
        return "-", None, None, None
    idx = tmp[metric].fillna(-1).idxmax()
    row = tmp.loc[idx]
    best_model = row["Model"] if "Model" in row else str(idx)
    return best_model, row.get("F1"), row.get("Accuracy"), row.get("AUC-ROC")


def make_run_id(config: Dict) -> str:
    base = f"{config['dataset']}_{int(time.time())}_{st.session_state.get('run_counter', 0)}"
    return hashlib.md5(base.encode()).hexdigest()[:12]


def collect_output_files(base_dir: Path, dataset: str) -> Tuple[List[Path], List[Path]]:
    fig_dir = base_dir / "outputs" / "figures"
    log_dir = base_dir / "outputs" / "logs"
    fig_paths = sorted(fig_dir.glob(f"*{dataset}*.png"), key=lambda p: p.stat().st_mtime, reverse=True) if fig_dir.exists() else []
    log_paths = sorted(log_dir.glob(f"*{dataset}*.log"), key=lambda p: p.stat().st_mtime, reverse=True) if log_dir.exists() else []
    return fig_paths, log_paths


def nice_figure_label(path: Path) -> str:
    return path.stem.replace("_", " ")


def make_methodology_manifest(result_bundle: Optional[Dict], config: Optional[Dict] = None, audit: Optional[Dict] = None, guards: Optional[Dict] = None) -> Dict:
    cfg = config or (result_bundle or {}).get("config", {}) or {}
    summary_df = (result_bundle or {}).get("result", {}).get("summary") if result_bundle else None
    best_model, best_f1, best_acc, best_auc = get_best_model_info(summary_df)
    ml_active, dl_active = active_model_groups(cfg)
    manifest = {
        "experiment_name": cfg.get("experiment_name", "run"),
        "dataset": cfg.get("dataset", "-"),
        "dataset_label": DATASET_LABELS.get(cfg.get("dataset", ""), cfg.get("dataset", "-")),
        "classification_mode": "Binary" if cfg.get("binary", True) else "Multiclass",
        "workflow_mode": cfg.get("workflow_mode", "exploratory"),
        "experiment_type": cfg.get("experiment_type", "custom"),
        "input_mode": cfg.get("input_mode", "Path manual"),
        "split_protocol": SPLIT_PROTOCOLS.get(cfg.get("dataset", ""), "Dataset-specific split protocol."),
        "scaler": cfg.get("resolved_scaler", cfg.get("scaler", "-")),
        "balancing": cfg.get("balance", "-"),
        "copula_family": cfg.get("resolved_copula", cfg.get("copula_family", "-")),
        "top_k_features": cfg.get("top_k"),
        "cv_folds": cfg.get("cv_folds"),
        "cicids_sample_fraction": cfg.get("sample_frac"),
        "cicids_split": cfg.get("cicids_split"),
        "seed": cfg.get("seed"),
        "active_ml_models": ml_active,
        "active_dl_models": dl_active,
        "bayesian_reasoning": bool(cfg.get("use_bn")),
        "shap_enabled": bool(cfg.get("use_shap")),
        "complexity_label": cfg.get("complexity_label"),
        "complexity_score": cfg.get("complexity_score"),
        "theme": THEME_OPTIONS.get(cfg.get("theme", "cyber"), THEME_OPTIONS["cyber"])["label"],
        "timestamp": (result_bundle or {}).get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        "notes": cfg.get("notes", ""),
        "protocol_lock": bool(cfg.get("protocol_lock")),
        "dataset_final_preset": cfg.get("dataset_final_preset", "None"),
        "dataset_final_preset_description": cfg.get("dataset_final_preset_description", ""),
        "report_title": cfg.get("report_title", "Hybrid NIDS Academic Report"),
        "report_author": cfg.get("report_author", "Researcher"),
        "report_affiliation": cfg.get("report_affiliation", "Graduate Research Program"),
        "report_template": cfg.get("report_template", "formal_dissertation"),
        "methodological_guards": guards or {},
        "dataset_audit": audit or {},
        "pipeline_path": cfg.get("pipeline_path"),
        "environment": {
            "python": sys.version.split()[0],
            "platform": platform.platform(),
            "streamlit": st.__version__,
            "pandas": pd.__version__,
        },
        "best_model": best_model,
        "best_f1": best_f1,
        "best_accuracy": best_acc,
        "best_auc": best_auc,
        "experimental_notes": [
            "AAE refers to an autoencoder-regularized classifier scaffold, not a full adversarial autoencoder.",
            "Transfer Learning refers to self-supervised encoder pretraining on the training split, then downstream fine-tuning.",
        ],
    }
    return manifest


def make_academic_report_markdown(result_bundle: Dict, config: Dict, audit: Optional[Dict], guards: Dict, style: str = "thesis") -> str:
    summary_df = result_bundle.get("result", {}).get("summary") if result_bundle else None
    best_model, best_f1, best_acc, best_auc = get_best_model_info(summary_df)
    dataset_label = DATASET_LABELS.get(config.get("dataset", ""), config.get("dataset", "-"))
    mode = "binary" if config.get("binary", True) else "multiclass"
    scaler = config.get("resolved_scaler", config.get("scaler", "-"))
    balance = config.get("balance", "-")
    copula = config.get("resolved_copula", config.get("copula_family", "-"))
    ml_active, dl_active = active_model_groups(config)
    split = SPLIT_PROTOCOLS.get(config.get("dataset", ""), "dataset-specific split protocol")
    readiness = guards.get("label", "-")
    audit_rows = audit.get("overall", {}).get("total_rows_estimate", "N/A") if audit else "N/A"

    f1_text = "N/A" if best_f1 is None or pd.isna(best_f1) else f"{float(best_f1):.4f}"
    acc_text = "N/A" if best_acc is None or pd.isna(best_acc) else f"{float(best_acc):.4f}"
    auc_text = "N/A" if best_auc is None or pd.isna(best_auc) else f"{float(best_auc):.4f}"

    if style == "formal_journal":
        intro = f"# {config.get('report_title', f'Hybrid NIDS Experiment on {dataset_label}')}\n\n**Author:** {config.get('report_author', 'Researcher')}  \n**Affiliation:** {config.get('report_affiliation', 'Graduate Research Program')}\n\nThis report summarises a research-grade Network Intrusion Detection System experiment on the {dataset_label} dataset. "
    elif style == "formal_dissertation":
        intro = f"# {config.get('report_title', f'Hybrid NIDS Experiment on {dataset_label}')}\n\n**Author:** {config.get('report_author', 'Researcher')}  \n**Affiliation:** {config.get('report_affiliation', 'Graduate Research Program')}\n\nPada eksperimen ini, dataset {dataset_label} digunakan untuk membangun pipeline Hybrid NIDS dalam mode {mode}. "
    elif style == "chapter4_results":
        intro = f"# BAB IV — HASIL DAN PEMBAHASAN\n\n## 4.1 Gambaran Umum Eksperimen\nEksperimen pada bagian ini menggunakan dataset **{dataset_label}** dalam mode **{mode}**. Protokol validasi yang digunakan adalah **{config.get('validation_method', 'holdout')}** dengan skema split **{split}**. Semua transformasi penting dipelajari dari data train sehingga evaluasi akhir tetap adil dan konsisten secara metodologis."
    elif style == "paper":
        intro = (
            f"This experiment evaluates a hybrid Network Intrusion Detection System on the {dataset_label} dataset using a {mode} classification setting. "
            f"The protocol preserves separation between training, validation, and final testing while combining ML and DL experts under a copula-based fusion design."
        )
    elif style == "concise":
        intro = f"Run summary for {dataset_label}: {mode} setup with {scaler} scaling, {balance} balancing, and {copula} fusion."
    else:
        intro = (
            f"Pada eksperimen ini, dataset {dataset_label} digunakan untuk membangun pipeline Hybrid NIDS dalam mode {mode}. "
            f"Protokol split yang dipakai adalah: {split} Semua transformasi penting dipelajari dari data train agar evaluasi akhir tetap adil dan dapat dipertanggungjawabkan secara metodologis."
        )

    body = (
        f"\n\n## 4.2 Konfigurasi Model dan Validasi\nKelompok model Machine Learning (ML) yang aktif adalah: {', '.join(ml_active) or 'tidak ada'}. "
        f"Kelompok model Deep Learning (DL) yang aktif adalah: {', '.join(dl_active) or 'tidak ada'}. "
        f"Konfigurasi preprocessing utama menggunakan scaler **{scaler}**, balancing **{balance}**, dan copula **{copula}**. "
        f"Metode seleksi fitur yang digunakan adalah **{config.get('feature_method', '-') }**, sedangkan skema validasi utama adalah **{config.get('validation_method', '-') }**. "
        f"Audit dataset memperkirakan total baris sekitar **{audit_rows}** dan guard metodologis saat ini berada pada status **{readiness}**. "
        f"\n\n## 4.3 Ringkasan Hasil Eksperimen\nModel terbaik pada ringkasan saat ini adalah **{best_model}** dengan F1 = **{f1_text}**, Accuracy = **{acc_text}**, dan AUC-ROC = **{auc_text}**. "
        f"Hasil ini perlu dibaca bersama confusion matrix, kurva ROC/PR, distribusi kelas, dan analisis korelasi fitur yang dihasilkan pipeline analytics. "
        f"\n\n## 4.4 Pembahasan Singkat\nSecara umum, kombinasi model ML dan DL aktif menunjukkan desain eksperimen yang mengarah pada pendekatan hybrid. Pemilihan scaler **{scaler}**, balancing **{balance}**, dan copula **{copula}** merefleksikan strategi untuk menjaga kestabilan performa sekaligus mengevaluasi fusi probabilistik antarmodel."
    )
    closing = (
        "\n\n## 4.5 Catatan Akademik\nHasil pada bagian ini sebaiknya diinterpretasikan bersama detail split dataset, strategi balancing, komposisi model ML/DL yang aktif, audit dataset, guard metodologis, serta manifest reproduksibilitas yang dihasilkan dashboard."
    )
    return intro + body + closing


def build_results_zip(result_bundle: Dict, audit: Optional[Dict], guards: Dict, manifest: Dict, report_md: str, report_docx: Optional[bytes] = None, report_pdf: Optional[bytes] = None) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        summary_df = result_bundle.get("result", {}).get("summary")
        if isinstance(summary_df, pd.DataFrame):
            zf.writestr(f"summary_{result_bundle['config']['dataset']}.csv", summary_df.to_csv(index=False))
        logs = result_bundle.get("logs", "")
        if logs:
            zf.writestr("live_logs.txt", logs)
        zf.writestr("config.json", json.dumps(result_bundle.get("config", {}), indent=2, default=str))
        zf.writestr("manifest.json", json.dumps(manifest, indent=2, default=str))
        zf.writestr("dataset_audit.json", json.dumps(audit or {}, indent=2, default=str))
        zf.writestr("guardrails.json", json.dumps(guards, indent=2, default=str))
        zf.writestr("report_note.md", report_md)
        if report_docx:
            zf.writestr("report.docx", report_docx)
        if report_pdf:
            zf.writestr("report.pdf", report_pdf)
        for fig in result_bundle.get("fig_paths", []):
            if Path(fig).exists():
                zf.write(fig, arcname=f"figures/{Path(fig).name}")
        for log in result_bundle.get("log_paths", [])[:5]:
            if Path(log).exists():
                zf.write(log, arcname=f"logs/{Path(log).name}")
    buffer.seek(0)
    return buffer.getvalue()


def save_artifacts(run_id: str, result_bundle: Dict, audit: Optional[Dict], guards: Dict, manifest: Dict, report_md: str) -> Tuple[Path, Path]:
    out_dir = ARTIFACT_ROOT / run_id
    out_dir.mkdir(parents=True, exist_ok=True)
    summary_df = result_bundle.get("result", {}).get("summary")
    if isinstance(summary_df, pd.DataFrame):
        summary_df.to_csv(out_dir / "summary.csv", index=False)
    (out_dir / "manifest.json").write_text(json.dumps(manifest, indent=2, default=str), encoding="utf-8")
    (out_dir / "dataset_audit.json").write_text(json.dumps(audit or {}, indent=2, default=str), encoding="utf-8")
    (out_dir / "guardrails.json").write_text(json.dumps(guards or {}, indent=2, default=str), encoding="utf-8")
    (out_dir / "report_note.md").write_text(report_md, encoding="utf-8")
    report_docx = None
    report_pdf = None
    try:
        report_docx = build_report_docx_bytes("Hybrid NIDS Academic Report", report_md, manifest=manifest, fig_paths=result_bundle.get("fig_paths", []))
        (out_dir / "report.docx").write_bytes(report_docx)
    except Exception:
        report_docx = None
    try:
        report_pdf = build_report_pdf_bytes("Hybrid NIDS Academic Report", report_md, manifest=manifest, fig_paths=result_bundle.get("fig_paths", []))
        (out_dir / "report.pdf").write_bytes(report_pdf)
    except Exception:
        report_pdf = None
    if result_bundle.get("logs"):
        (out_dir / "live_logs.txt").write_text(result_bundle["logs"], encoding="utf-8")
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(exist_ok=True)
    for fig in result_bundle.get("fig_paths", []):
        fig = Path(fig)
        if fig.exists():
            try:
                shutil.copy2(fig, fig_dir / fig.name)
            except Exception:
                pass
    log_dir = out_dir / "logs"
    log_dir.mkdir(exist_ok=True)
    for log in result_bundle.get("log_paths", [])[:20]:
        log = Path(log)
        if log.exists():
            try:
                shutil.copy2(log, log_dir / log.name)
            except Exception:
                pass
    preprep_report = result_bundle.get("result", {}).get("prepreprocess_report")
    if preprep_report:
        (out_dir / "prepreprocess_report.json").write_text(json.dumps(preprep_report, indent=2, default=str), encoding="utf-8")
    zip_bytes = build_results_zip(result_bundle, audit, guards, manifest, report_md, report_docx=report_docx, report_pdf=report_pdf)
    zip_path = out_dir / f"{run_id}_academic_package.zip"
    zip_path.write_bytes(zip_bytes)
    return out_dir, zip_path

def db_insert_running(run_id: str, config: Dict, audit: Optional[Dict], guards: Dict):
    ml_active, dl_active = active_model_groups(config)
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            INSERT OR REPLACE INTO experiments (
                run_id, created_at, updated_at, status, experiment_name, dataset, classification_mode, workflow_mode, experiment_type,
                scaler, balance, copula, cv_folds, top_k, seed, complexity_score, complexity_label, notes, ml_models, dl_models,
                theme, protocol_locked, pipeline_path, config_json, audit_json, guards_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                run_id, created_at, created_at, "running", config.get("experiment_name"), config.get("dataset"),
                "Binary" if config.get("binary", True) else "Multiclass", config.get("workflow_mode"), config.get("experiment_type"),
                config.get("resolved_scaler", config.get("scaler")), config.get("balance"), config.get("resolved_copula", config.get("copula_family")),
                int(config.get("cv_folds", 0)), int(config.get("top_k", 0)), int(config.get("seed", 0)),
                int(config.get("complexity_score", 0)), config.get("complexity_label"), config.get("notes", ""),
                ", ".join(ml_active), ", ".join(dl_active), THEME_OPTIONS.get(config.get("theme", "cyber"), THEME_OPTIONS["cyber"])["label"],
                1 if config.get("protocol_lock") else 0, config.get("pipeline_path"), json.dumps(config, default=str), json.dumps(audit or {}, default=str), json.dumps(guards, default=str)
            )
        )
        conn.commit()


def db_complete_run(run_id: str, status: str, config: Dict, manifest: Optional[Dict] = None, audit: Optional[Dict] = None, guards: Optional[Dict] = None,
                    result_bundle: Optional[Dict] = None, artifact_zip_path: Optional[Path] = None, output_root: Optional[Path] = None, error_text: str = ""):
    summary_df = result_bundle.get("result", {}).get("summary") if result_bundle else None
    best_model, best_f1, best_acc, best_auc = get_best_model_info(summary_df)
    updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            UPDATE experiments
            SET updated_at=?, status=?, best_model=?, f1=?, accuracy=?, auc=?, manifest_json=?, audit_json=?, guards_json=?, artifact_zip_path=?, output_root=?, error_text=?
            WHERE run_id=?
            """,
            (
                updated_at, status, best_model,
                None if best_f1 is None else float(best_f1),
                None if best_acc is None else float(best_acc),
                None if best_auc is None else float(best_auc),
                json.dumps(manifest or {}, default=str), json.dumps(audit or {}, default=str), json.dumps(guards or {}, default=str),
                str(artifact_zip_path) if artifact_zip_path else None, str(output_root) if output_root else None, error_text, run_id,
            )
        )
        conn.commit()




def strip_markdown_for_export(md: str) -> str:
    lines = []
    for raw in (md or "").splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        while line.startswith("#"):
            line = line[1:].strip()
        line = line.replace("**", "").replace("__", "").replace("`", "")
        line = line.replace("- ", "• ")
        lines.append(line)
    return "\n".join(lines).strip()


def select_recommended_figures(fig_paths: List[Path], max_items: int = 6) -> List[Path]:
    if not fig_paths:
        return []
    preferred_keywords = [
        "comparison", "summary_heatmap", "roc", "pr_curve", "class_distribution",
        "correlation_heatmap", "confusion", "fused_dist", "boxplots"
    ]
    chosen: List[Path] = []
    seen = set()
    for kw in preferred_keywords:
        for p in fig_paths:
            key = str(p)
            if key in seen:
                continue
            if kw in p.name.lower():
                chosen.append(p)
                seen.add(key)
                break
        if len(chosen) >= max_items:
            return chosen[:max_items]
    for p in fig_paths:
        key = str(p)
        if key in seen:
            continue
        chosen.append(p)
        seen.add(key)
        if len(chosen) >= max_items:
            break
    return chosen[:max_items]


def figure_label_map(fig_paths: List[Path]) -> Dict[str, Path]:
    return {nice_figure_label(Path(p)): Path(p) for p in fig_paths}


def build_report_docx_bytes(title: str, markdown_text: str, manifest: Optional[Dict] = None, fig_paths: Optional[List[Path]] = None) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx belum terpasang. Install: pip install python-docx")
    doc = Document()
    doc.add_heading(title, 0)
    if manifest:
        p = doc.add_paragraph()
        p.add_run(f"Dataset: {manifest.get('dataset_label', manifest.get('dataset', '-'))}\n")
        p.add_run(f"Mode: {manifest.get('classification_mode', '-')}\n")
        p.add_run(f"Workflow: {manifest.get('workflow_mode', '-')}\n")
        p.add_run(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for line in strip_markdown_for_export(markdown_text).splitlines():
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("• "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    selected_figs = [Path(p) for p in (fig_paths or []) if Path(p).exists()][:6]
    if selected_figs and Inches is not None:
        doc.add_page_break()
        doc.add_heading("Selected analytics figures", level=1)
        for fig in selected_figs:
            doc.add_paragraph(nice_figure_label(fig))
            try:
                doc.add_picture(str(fig), width=Inches(6.2))
            except Exception:
                doc.add_paragraph(f"[Figure unavailable: {fig.name}]")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_report_pdf_bytes(title: str, markdown_text: str, manifest: Optional[Dict] = None, fig_paths: Optional[List[Path]] = None) -> bytes:
    if canvas is None or A4 is None:
        raise RuntimeError("reportlab belum terpasang. Install: pip install reportlab")
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    x, y = 50, height - 50
    c.setFont('Helvetica-Bold', 14)
    c.drawString(x, y, title)
    y -= 22
    c.setFont('Helvetica', 10)
    manifest_lines = []
    if manifest:
        manifest_lines = [
            f"Dataset: {manifest.get('dataset_label', manifest.get('dataset', '-'))}",
            f"Mode: {manifest.get('classification_mode', '-')}",
            f"Workflow: {manifest.get('workflow_mode', '-')}",
            f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
        ]
    text_lines = manifest_lines + strip_markdown_for_export(markdown_text).splitlines()
    for line in text_lines:
        if y < 50:
            c.showPage()
            c.setFont('Helvetica', 10)
            y = height - 50
        c.drawString(x, y, str(line)[:120])
        y -= 14
    selected_figs = [Path(p) for p in (fig_paths or []) if Path(p).exists()][:4]
    if selected_figs and ImageReader is not None:
        for fig in selected_figs:
            c.showPage()
            y = height - 40
            c.setFont('Helvetica-Bold', 12)
            c.drawString(40, y, nice_figure_label(fig))
            y -= 20
            try:
                img = ImageReader(str(fig))
                iw, ih = img.getSize()
                max_w = width - 80
                max_h = height - 120
                scale = min(max_w / max(iw, 1), max_h / max(ih, 1))
                draw_w = iw * scale
                draw_h = ih * scale
                c.drawImage(img, 40, max(40, y - draw_h), width=draw_w, height=draw_h, preserveAspectRatio=True, mask='auto')
            except Exception:
                c.setFont('Helvetica', 10)
                c.drawString(40, y, f"[Figure unavailable: {fig.name}]")
    c.save()
    return bio.getvalue()


def infer_job_progress(log_text: str, status: str = '') -> Dict:
    text = (log_text or '').lower()
    if status == 'queued':
        return {'label': 'Queued', 'value': 5, 'stage': 'Queue'}
    if status == 'stopped':
        return {'label': 'Stopped', 'value': 0, 'stage': 'Stopped'}
    if status == 'failed':
        return {'label': 'Failed', 'value': 100, 'stage': 'Failed'}
    if status == 'success':
        return {'label': 'Completed', 'value': 100, 'stage': 'Done'}
    stage_map = [
        ('stage 1', 10, 'Loading'), ('stage 2', 20, 'EDA'), ('stage 3', 35, 'Preprocessing'),
        ('stage 4', 55, 'Model training'), ('stage 5', 68, 'Copula fusion'), ('stage 6', 76, 'Distribution analysis'),
        ('stage 7', 86, 'Bayesian reasoning'), ('stage 8', 94, 'Evaluation'), ('pipeline done', 100, 'Done'),
    ]
    for marker, pct, stage in reversed(stage_map):
        if marker in text:
            return {'label': f'{stage} ({pct}%)', 'value': pct, 'stage': stage}
    if 'training' in text:
        return {'label': 'Training started (45%)', 'value': 45, 'stage': 'Training'}
    if 'running' in status:
        return {'label': 'Running (12%)', 'value': 12, 'stage': 'Initializing'}
    return {'label': 'Unknown', 'value': 0, 'stage': 'Unknown'}


def read_job_live_log(job_dir: Optional[Path]) -> str:
    if not job_dir:
        return ''
    p = Path(job_dir) / 'worker_live.log'
    if p.exists():
        return p.read_text(encoding='utf-8', errors='ignore')
    p2 = Path(job_dir) / 'worker_error.txt'
    return p2.read_text(encoding='utf-8', errors='ignore') if p2.exists() else ''

def load_registry_df() -> pd.DataFrame:
    with sqlite3.connect(DB_PATH) as conn:
        df = pd.read_sql_query("SELECT * FROM experiments ORDER BY created_at DESC", conn)
    return df


def make_history_row(run_id: str, result_bundle: Dict, manifest: Dict) -> Dict:
    summary_df = result_bundle.get("result", {}).get("summary") if result_bundle else None
    best_model, best_f1, best_acc, best_auc = get_best_model_info(summary_df)
    return {
        "run_id": run_id,
        "timestamp": manifest.get("timestamp"),
        "experiment_name": manifest.get("experiment_name"),
        "dataset": manifest.get("dataset_label"),
        "mode": manifest.get("classification_mode"),
        "workflow": manifest.get("workflow_mode"),
        "experiment_type": manifest.get("experiment_type"),
        "scaler": manifest.get("scaler"),
        "balance": manifest.get("balancing"),
        "copula": manifest.get("copula_family"),
        "best_model": best_model,
        "F1": best_f1,
        "Accuracy": best_acc,
        "AUC-ROC": best_auc,
    }


def run_pipeline(module, config: Dict, log_placeholder):
    pipeline_cls = getattr(module, "NIDSPipeline")
    writer = StreamlitLogWriter(log_placeholder)
    run_args = dict(
        dataset=config["dataset"], paths=config["paths"], binary=config["binary"], copula_family=config["copula_family"],
        balance=config["balance"], scaler=config["scaler"], top_k=config["top_k"], cv_folds=config["cv_folds"],
        use_lr=config.get("use_lr", True), use_rf=config.get("use_rf", True), use_svm=config.get("use_svm", True), use_xgb=config.get("use_xgb", True), use_dt=config.get("use_dt", True),
        use_mlp=config.get("use_mlp", True), use_dnn=config.get("use_dnn", False), use_cnn1d=config.get("use_cnn1d", False), use_lstm=config.get("use_lstm", False),
        use_aae=config.get("use_aae", False), use_cnn_lstm=config.get("use_cnn_lstm", False), use_cnn_bilstm=config.get("use_cnn_bilstm", True), use_cnn_gru=config.get("use_cnn_gru", True),
        use_gnn=config.get("use_gnn", False), use_gat=config.get("use_gat", False), use_transformer=config.get("use_transformer", True), use_transfer_learning=config.get("use_transfer_learning", False),
        use_bn=config["use_bn"], use_shap=config["use_shap"], sample_frac=config["sample_frac"], cicids_split=config["cicids_split"],
        model_params=config.get("model_params", {}), training_params=config.get("training_params", {}),
        learning_paradigm=config.get("learning_paradigm", "supervised"),
        use_label_spreading=config.get("use_label_spreading", False), use_self_training_svm=config.get("use_self_training_svm", False),
        use_self_training_rf=config.get("use_self_training_rf", False), use_pseudo_label_dnn=config.get("use_pseudo_label_dnn", False),
        use_isolation_forest=config.get("use_isolation_forest", False), use_oneclass_svm_anom=config.get("use_oneclass_svm_anom", False),
        use_lof=config.get("use_lof", False), use_pca_anom=config.get("use_pca_anom", False), use_kmeans_anom=config.get("use_kmeans_anom", False),
        use_autoencoder_anom=config.get("use_autoencoder_anom", False), use_rl_threshold=config.get("use_rl_threshold", False),
        semi_label_fraction=config.get("semi_label_fraction", 0.3), pseudo_label_threshold=config.get("pseudo_label_threshold", 0.9),
        anomaly_contamination=config.get("anomaly_contamination", 0.1), rl_episodes=config.get("rl_episodes", 40),
    )
    pipeline = pipeline_cls(**run_args)
    start = time.time()
    with contextlib.redirect_stdout(writer), contextlib.redirect_stderr(writer):
        result = pipeline.run()
    writer.flush()
    elapsed = time.time() - start
    return result, writer.value, elapsed


def figure_category(path: Path) -> str:
    name = path.stem.lower()
    if name.startswith("eda_") or "class_distribution" in name or "correlation_heatmap" in name or "pairplot" in name or "attack_breakdown" in name or "missing_values" in name or "prepreprocess_audit" in name or "feature_label_heatmap" in name:
        return "EDA & Data Quality"
    if name.startswith("cm_") or name.startswith("roc_") or name.startswith("pr_curve_") or "comparison" in name or "summary_heatmap" in name or "fused_dist" in name:
        return "Evaluation & Model Comparison"
    return "Other"


def render_prepreprocessing_report(result_bundle: Optional[Dict]):
    st.markdown("### 🧼 Pre-preprocessing & preprocessing audit")
    report = (result_bundle or {}).get("result", {}).get("prepreprocess_report")
    if not report:
        st.info("Laporan pre-preprocessing belum tersedia. Jalankan pipeline analytics atau buka artifact run terbaru.")
        return
    before = report.get("before", {})
    after = report.get("after", {})
    actions = report.get("actions", {})
    c1, c2, c3, c4 = st.columns(4)
    metric_card("Rows before", str(before.get("rows", "N/A")), "Ukuran train split sebelum hygiene stage")
    metric_card("Rows after", str(after.get("rows", "N/A")), "Ukuran train split setelah hygiene stage")
    metric_card("Duplicates removed", str(actions.get("duplicates_removed", 0)), "Row duplicate yang dihapus pada stage awal")
    metric_card("Constant features dropped", str(actions.get("constant_features_dropped", 0)), "Fitur konstan yang dihapus")

    c5, c6, c7 = st.columns(3)
    mini_card("Missing % before", f"{float(before.get('missing_pct', 0) or 0):.2f}%", "Rasio missing value")
    mini_card("Missing % after", f"{float(after.get('missing_pct', 0) or 0):.2f}%", "Setelah pre-preprocessing")
    mini_card("Infinite values fixed", str(actions.get("inf_values_replaced", 0)), "Jumlah inf/-inf yang disanitasi")

    dropped = report.get("dropped_constant_features", []) or []
    if dropped:
        with st.expander("Daftar constant features yang dihapus", expanded=False):
            st.dataframe(pd.DataFrame({"constant_feature": dropped}), use_container_width=True, hide_index=True)

    with st.expander("Raw pre-preprocessing report (JSON)", expanded=False):
        st.json(report)


def render_analytics_gallery(result_bundle: Optional[Dict], key_prefix: str = "analytics_gallery"):
    st.markdown("### 📊 Analytics gallery")
    fig_paths = [Path(p) for p in (result_bundle or {}).get("fig_paths", [])]
    if not fig_paths:
        st.info("Belum ada figure analytics yang tersedia.")
        return
    categories: Dict[str, List[Path]] = {}
    for path in fig_paths:
        categories.setdefault(figure_category(path), []).append(path)
    tabs = st.tabs(list(categories.keys()))
    thumb_cols = st.slider("Jumlah kolom thumbnail", min_value=2, max_value=5, value=3, key=f"{key_prefix}_thumbcols")
    for tab, cat in zip(tabs, categories.keys()):
        with tab:
            paths = categories[cat]
            label_map = figure_label_map(paths)
            default_choice = next(iter(label_map.keys())) if label_map else None
            choice = st.selectbox(f"Pilih figure — {cat}", list(label_map.keys()), key=f"{key_prefix}_{cat}") if label_map else None
            if choice:
                st.image(str(label_map[choice]), use_container_width=True, caption=choice)
            st.caption(f"{len(paths)} figure dalam kategori {cat}.")
            with st.expander("Thumbnail gallery", expanded=False):
                cols = st.columns(thumb_cols)
                for i, p in enumerate(paths):
                    with cols[i % thumb_cols]:
                        st.image(str(p), use_container_width=True)
                        st.caption(nice_figure_label(p))


def render_dataset_audit(audit: Optional[Dict]):
    st.markdown("### 🔎 Dataset audit report")
    if not audit:
        st.info("Audit dataset belum tersedia. Pastikan path valid atau upload selesai.")
        return
    overall = audit.get("overall", {})
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        metric_card("Files detected", str(overall.get("n_files", 0)), "Jumlah file sumber yang berhasil dibaca")
    with c2:
        metric_card("Rows estimate", str(overall.get("total_rows_estimate", "N/A")), "Perkiraan total baris berdasarkan line count")
    with c3:
        metric_card("Sample missing %", f"{overall.get('sample_missing_pct', 0) or 0:.2f}%", "Rasio missing pada sample audit")
    with c4:
        metric_card("Sample duplicate %", f"{overall.get('sample_duplicate_pct', 0) or 0:.2f}%", "Rasio duplicate pada sample audit")

    files_df = pd.DataFrame(audit.get("files", []))
    if not files_df.empty:
        st.dataframe(files_df, use_container_width=True, hide_index=True)


def render_guardrails(guards: Dict):
    st.markdown("### 🧷 Leakage guard & methodological checks")
    if not guards:
        st.info("Guardrails belum dihitung.")
        return
    st.markdown(f"**Guardrail score:** {guards.get('label', '-')} ({guards.get('score', '-')}/100)")
    st.progress(guards.get("score", 0) / 100)
    for idx, item in enumerate(guards.get("checks", [])):
        css = "audit-ok" if item["status"] == "pass" else "audit-warn" if item["status"] == "warn" else "audit-bad"
        icon = "✅" if item["status"] == "pass" else "⚠️" if item["status"] == "warn" else "⛔"
        st.markdown(f"<div class='glass'><div class='{css}'>{icon} {item['title']}</div><div class='small-muted'>{item['detail']}</div></div>", unsafe_allow_html=True)


def render_hero(config: Dict, history_count: int, readiness_score: int, readiness_label: str):
    ml_active, dl_active = active_model_groups(config)
    pills = [
        f"Dataset · {DATASET_LABELS.get(config['dataset'], config['dataset'])}",
        f"Mode · {'Binary' if config.get('binary', True) else 'Multiclass'}",
        f"Workflow · {WORKFLOW_MODES.get(config.get('workflow_mode'), config.get('workflow_mode'))}",
        f"Experiment · {EXPERIMENT_TYPES.get(config.get('experiment_type'), config.get('experiment_type'))}",
        f"Scaler · {config.get('resolved_scaler')}",
        f"Copula · {config.get('resolved_copula')}",
    ]
    st.markdown(
        f"""
        <div class='hero'>
            <h1>🛡️ Hybrid NIDS Research Studio v4.6.1 Final Study</h1>
            <p>Dashboard eksperimen yang lebih research-grade dengan pre-preprocessing audit, enhanced EDA, dan evaluasi visual lengkap.</p>
            <div>{''.join(f"<span class='pill'>{p}</span>" for p in pills)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        metric_card("Readiness", f"{readiness_label} ({readiness_score}/100)", "Semakin tinggi, semakin siap untuk dijalankan")
    with c2:
        metric_card("ML models aktif", str(len(ml_active)), ", ".join(ml_active) or "Belum ada")
    with c3:
        metric_card("DL models aktif", str(len(dl_active)), ", ".join(dl_active) or "Belum ada")
    with c4:
        metric_card("Session history", str(history_count), "Run yang tersimpan di session saat ini")


def render_overview_cards(latest: Optional[Dict], config: Dict, guards: Dict, audit: Optional[Dict]):
    c1, c2, c3, c4 = st.columns(4)
    best_model, best_f1, best_acc, best_auc = get_best_model_info((latest or {}).get("result", {}).get("summary"))
    with c1:
        metric_card("Workflow", WORKFLOW_MODES.get(config.get("workflow_mode"), config.get("workflow_mode")), "Gunakan publication mode untuk run yang akan dilaporkan")
    with c2:
        metric_card("Complexity", f"{config.get('complexity_label')} ({config.get('complexity_score')})", config.get("complexity_note", ""))
    with c3:
        metric_card("Guard score", f"{guards.get('label')} ({guards.get('score')}/100)", "Audit metodologi sebelum run")
    with c4:
        metric_card("Best current F1", "N/A" if best_f1 is None else f"{float(best_f1):.4f}", f"Model: {best_model}")

    c5, c6, c7, c8 = st.columns(4)
    with c5:
        mini_card("Rows estimate", str((audit or {}).get("overall", {}).get("total_rows_estimate", "N/A")), "Perkiraan total baris dataset")
    with c6:
        mini_card("Split protocol", SPLIT_PROTOCOLS.get(config["dataset"], "Dataset specific"), "Rancangan train/validation/test")
    with c7:
        mini_card("Protocol lock", "Locked" if config.get("protocol_lock") else "Unlocked", "Run final sebaiknya selalu locked")
    with c8:
        mini_card("Registry persistence", str(len(load_registry_df())), "Jumlah total run di SQLite registry")


def render_config_panel(config: Dict, readiness_score: int, readiness_label: str, readiness_notes: List[str]):
    ml_active, dl_active = active_model_groups(config)
    left, right = st.columns([1.25, 1])
    with left:
        st.markdown("### 🧭 Experiment blueprint")
        st.markdown(
            f"<div class='glass'><div class='metric-label'>Research protocol</div>"
            f"<div class='metric-value' style='font-size:1.05rem'>{EXPERIMENT_TYPES.get(config['experiment_type'])}</div>"
            f"<div class='small-muted'>Dataset: <strong>{DATASET_LABELS[config['dataset']]}</strong> · Mode: <strong>{'Binary' if config['binary'] else 'Multiclass'}</strong> · Workflow: <strong>{WORKFLOW_MODES[config['workflow_mode']]}</strong><br>Paradigm: <strong>{LEARNING_PARADIGM_LABELS.get(config.get('learning_paradigm', 'supervised'), config.get('learning_paradigm', 'supervised'))}</strong> · Track: <strong>{SUPERVISED_TRACK_LABELS.get(config.get('supervised_track', 'both'), config.get('supervised_track', 'both')) if config.get('learning_paradigm', 'supervised') == 'supervised' else 'Backend aktif / experimental'}</strong><br>Scaler: <strong>{config['resolved_scaler']}</strong> · Balancing: <strong>{config['balance']}</strong> · Copula: <strong>{config['resolved_copula']}</strong><br>CV: <strong>{config['cv_folds']}</strong> · Top-k: <strong>{config['top_k']}</strong> · Seed: <strong>{config['seed']}</strong></div></div>",
            unsafe_allow_html=True,
        )
        st.markdown("#### Model composition")
        st.caption(
            f"Paradigm utama: {LEARNING_PARADIGM_LABELS.get(config.get('learning_paradigm', 'supervised'), config.get('learning_paradigm', 'supervised'))}"
            + (f" · Track: {SUPERVISED_TRACK_LABELS.get(config.get('supervised_track', 'both'), config.get('supervised_track', 'both'))}" if config.get('learning_paradigm', 'supervised') == 'supervised' else " · Status: backend aktif / experimental")
        )
        c1, c2 = st.columns(2)
        with c1:
            mini_card("ML models", ", ".join(ml_active) or "Tidak ada", "Pastikan ada baseline klasik untuk pembanding")
        with c2:
            mini_card("DL models", ", ".join(dl_active) or "Tidak ada", "Aktifkan DL sesuai kapasitas komputasi")
    with right:
        st.markdown("### ✅ Readiness & config export")
        css = "readiness-good" if readiness_score >= 88 else "readiness-mid" if readiness_score >= 70 else "readiness-bad"
        st.markdown(f"<div class='glass'><div class='{css}'>Status: {readiness_label}</div><div class='small-muted'>Skor {readiness_score}/100</div></div>", unsafe_allow_html=True)
        for item in readiness_notes:
            st.caption(f"• {item}")
        st.download_button(
            "⬇️ Download current config JSON",
            data=json.dumps(config, indent=2, default=str).encode("utf-8"),
            file_name=f"config_{config['experiment_name']}.json",
            mime="application/json",
            key=f"config_json_{config['experiment_name']}_v2",
        )


def render_summary(result_bundle: Dict, key_prefix: str):
    result = result_bundle.get("result", {})
    summary_df = result.get("summary")
    feat_names = result.get("feat_names", [])
    fused_test = result.get("fused_test")
    tab1, tab2, tab3 = st.tabs(["Summary", "Selected features", "Fusion preview"])
    with tab1:
        if isinstance(summary_df, pd.DataFrame):
            st.dataframe(summary_df, use_container_width=True, hide_index=True)
            csv_bytes = summary_df.to_csv(index=False).encode("utf-8")
            st.download_button("⬇️ Download summary CSV", data=csv_bytes, file_name=f"results_{result_bundle['config']['dataset']}.csv", mime="text/csv", key=f"{key_prefix}_summary_csv")
        else:
            st.info("Summary belum tersedia.")
    with tab2:
        if feat_names:
            st.write(pd.DataFrame({"selected_feature": feat_names}))
        else:
            st.info("Daftar feature terpilih belum tersedia.")
    with tab3:
        if fused_test is not None:
            preview_df = pd.DataFrame(fused_test[:20], columns=[f"class_{i}" for i in range(fused_test.shape[1])])
            st.dataframe(preview_df, use_container_width=True, hide_index=True)
        else:
            st.info("Fusion preview belum tersedia.")


def render_figures(result_bundle: Dict, key_prefix: str):
    fig_paths = [Path(p) for p in result_bundle.get("fig_paths", [])]
    st.markdown("### 🖼️ Figures")
    if not fig_paths:
        st.info("Belum ada figure yang terdeteksi.")
        return
    recommended = select_recommended_figures(fig_paths, max_items=8)
    with st.expander("Recommended figures", expanded=True):
        cols = st.columns(min(4, max(1, len(recommended))))
        for i, p in enumerate(recommended):
            with cols[i % len(cols)]:
                st.image(str(p), use_container_width=True)
                st.caption(nice_figure_label(p))
    render_analytics_gallery(result_bundle, key_prefix=f"{key_prefix}_all")


def render_logs(result_bundle: Dict, key_prefix: str):
    logs = result_bundle.get("logs", "")
    log_paths = result_bundle.get("log_paths", [])
    st.markdown("### 📝 Logs")
    if logs:
        st.code(logs, language="text")
    else:
        st.info("Live log belum tersedia.")
    if log_paths:
        latest_log = log_paths[0]
        st.download_button("⬇️ Download latest pipeline log", data=latest_log.read_bytes(), file_name=latest_log.name, mime="text/plain", key=f"{key_prefix}_latest_log")


def render_study_design(config: Dict, latest: Optional[Dict], audit: Optional[Dict], guards: Dict):
    manifest = make_methodology_manifest(latest, config=config, audit=audit, guards=guards)
    ml_active, dl_active = manifest.get("active_ml_models", []), manifest.get("active_dl_models", [])
    st.markdown("### 🧪 Study design & methodological safeguards")
    c1, c2 = st.columns([1.15, 1])
    with c1:
        st.markdown(
            f"<div class='glass'><div style='font-size:1.05rem; font-weight:800; margin-bottom:0.4rem;'>Protocol summary</div><div class='small-muted'>Dataset: <strong>{manifest['dataset_label']}</strong> · Mode: <strong>{manifest['classification_mode']}</strong><br>Workflow: <strong>{manifest['workflow_mode']}</strong> · Experiment type: <strong>{manifest['experiment_type']}</strong><br>Split: <strong>{manifest['split_protocol']}</strong><br>Preprocessing: <strong>train-only fit</strong> · Resampling: <strong>train-only</strong> · Final evaluation: <strong>hold-out test</strong></div></div>",
            unsafe_allow_html=True,
        )
        st.markdown("#### Academic checklist")
        for title, detail in ACADEMIC_GUARDS:
            st.markdown(f"- **{title}** — {detail}")
    with c2:
        mini_card("ML family", ", ".join(ml_active) or "No active ML model")
        mini_card("DL family", ", ".join(dl_active) or "No active DL model")
        mini_card("Fusion design", f"Copula: {manifest['copula_family']} · Bayesian reasoning: {'Yes' if manifest['bayesian_reasoning'] else 'No'}")
        mini_card("Reproducibility", f"Seed: {manifest.get('seed')} · Theme: {manifest.get('theme')}")
    st.markdown("#### Guardrail findings")
    render_guardrails(guards)


def render_reproducibility_tab(config: Dict, latest: Optional[Dict], audit: Optional[Dict], guards: Dict, key_prefix: str):
    st.markdown("### ♻️ Reproducibility manifest")
    manifest = make_methodology_manifest(latest, config=config, audit=audit, guards=guards)
    st.json(manifest, expanded=False)
    manifest_bytes = json.dumps(manifest, indent=2, default=str).encode("utf-8")
    st.download_button("⬇️ Download manifest JSON", data=manifest_bytes, file_name=f"manifest_{manifest['dataset']}_{manifest['classification_mode'].lower()}.json", mime="application/json", key=f"{key_prefix}_manifest_json")


def render_report_tab(latest: Optional[Dict], config: Dict, audit: Optional[Dict], guards: Dict, key_prefix: str):
    st.markdown("### 📝 Academic report assistant")
    if not latest:
        st.info("Belum ada hasil run. Jalankan eksperimen terlebih dahulu agar report note bisa dibuat dari hasil aktual.")
        return
    style = st.selectbox("Report style", options=["formal_dissertation", "formal_journal", "chapter4_results", "concise", "thesis", "paper"], key=f"{key_prefix}_report_style")
    report_md = make_academic_report_markdown(latest, config=config, audit=audit, guards=guards, style=style)
    manifest = make_methodology_manifest(latest, config=config, audit=audit, guards=guards)
    st.markdown(f"**Title:** {config.get('report_title', 'Formal Hybrid NIDS Academic Report')}  ")
    st.markdown(f"**Author:** {config.get('report_author', 'Researcher')}  ")
    st.markdown(f"**Affiliation:** {config.get('report_affiliation', 'Graduate Research Program')}")
    st.markdown(report_md)
    st.code(report_md, language="markdown")
    c1, c2, c3 = st.columns(3)
    c1.download_button("⬇️ Download report note (.md)", data=report_md.encode("utf-8"), file_name=f"report_note_{latest['config']['dataset']}.md", mime="text/markdown", key=f"{key_prefix}_report_md")
    try:
        docx_bytes = build_report_docx_bytes("Hybrid NIDS Academic Report", report_md, manifest=manifest)
        c2.download_button("⬇️ Download report (.docx)", data=docx_bytes, file_name=f"report_{latest['config']['dataset']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{key_prefix}_report_docx")
    except Exception as exc:
        c2.info(f"DOCX export unavailable: {exc}")
    try:
        pdf_bytes = build_report_pdf_bytes("Hybrid NIDS Academic Report", report_md, manifest=manifest)
        c3.download_button("⬇️ Download report (.pdf)", data=pdf_bytes, file_name=f"report_{latest['config']['dataset']}.pdf", mime="application/pdf", key=f"{key_prefix}_report_pdf")
    except Exception as exc:
        c3.info(f"PDF export unavailable: {exc}")



def render_jobs_tab(key_prefix: str = "jobs"):
    st.markdown("### ⏱️ Background jobs & queue monitor")
    top = st.columns([1, 1, 1.2, 2.2])
    if top[0].button("🔄 Refresh jobs", key=f"{key_prefix}_refresh"):
        st.rerun()
    auto_flag = top[1].checkbox("Auto-refresh", value=st.session_state.get("jobs_auto_refresh", True), key=f"{key_prefix}_auto")
    st.session_state["jobs_auto_refresh"] = auto_flag
    refresh_sec = top[2].selectbox("Interval", options=[5, 10, 15, 30, 60], index=1, key=f"{key_prefix}_interval")
    top[3].caption("Progress diperkirakan dari state worker + stage log pipeline. Ini adalah monitor praktis, bukan persentase training yang presisi per epoch.")
    df = load_registry_df()
    if df.empty:
        st.info("Belum ada job yang tercatat.")
        return
    active = df[df["status"].isin(["queued", "running", "stopped", "failed", "success"])].copy()
    active = active.sort_values(["created_at"], ascending=[False]).head(30)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        metric_card("Queued", str(int((active["status"] == "queued").sum())), "Menunggu slot worker")
    with c2:
        metric_card("Running", str(int((active["status"] == "running").sum())), "Sedang diproses di background")
    with c3:
        metric_card("Succeeded", str(int((active["status"] == "success").sum())), "Run siap dimuat ulang")
    with c4:
        metric_card("Failed/Stopped", str(int((active["status"].isin(["failed", "stopped"])).sum())), "Perlu diagnosis atau resume")

    for _, row in active.iterrows():
        run_id = str(row.get("run_id"))
        status = str(row.get("status"))
        job_dir = Path(row.get("job_dir")) if row.get("job_dir") else None
        live_log = read_job_live_log(job_dir)
        prog = infer_job_progress(live_log, status=status)
        status_icon = {"queued": "🕓", "running": "🟢", "success": "✅", "failed": "❌", "stopped": "⏹️"}.get(status, "•")
        with st.container():
            cols = st.columns([3.3, 1.25, 1.1, 1.1, 1.0, 1.0])
            cols[0].markdown(f"<div class='glass'><strong>{status_icon} {row.get('experiment_name')}</strong><br><span class='small-muted'>run_id={run_id} · dataset={row.get('dataset')} · status={status} · updated={row.get('updated_at')}<br>Stage: <strong>{prog['stage']}</strong> · {prog['label']}</span></div>", unsafe_allow_html=True)
            cols[1].progress(max(0.0, min(prog['value'] / 100.0, 1.0)), text=f"{prog['value']}%")
            if status == "running":
                if cols[2].button("Stop", key=f"{key_prefix}_stop_{run_id}"):
                    stop_background_job(run_id)
                    st.rerun()
            else:
                cols[2].write("")
            if cols[3].button("Resume", key=f"{key_prefix}_resume_{run_id}"):
                try:
                    cfg = json.loads(row.get("config_json") or "{}")
                    aud = json.loads(row.get("audit_json") or "{}")
                    grd = json.loads(row.get("guards_json") or "{}")
                    new_id = enqueue_background_job(cfg, aud, grd)
                    start_next_queued_job()
                    st.success(f"Run baru diantrikan: {new_id}")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Gagal me-resume job: {exc}")
            if status == "success":
                if cols[4].button("Load", key=f"{key_prefix}_load_{run_id}"):
                    bundle = hydrate_result_bundle_from_row(row)
                    if bundle:
                        st.session_state["result_bundle"] = bundle
                        st.session_state["selected_registry_run"] = run_id
                        st.rerun()
            else:
                cols[4].write("")
            zip_path = row.get("artifact_zip_path")
            if zip_path and Path(str(zip_path)).exists():
                cols[5].download_button("ZIP", data=Path(str(zip_path)).read_bytes(), file_name=Path(str(zip_path)).name, mime="application/zip", key=f"{key_prefix}_zip_{run_id}")
            else:
                cols[5].write("")
            with st.expander(f"Live diagnostics — {run_id}", expanded=False):
                st.caption(f"Worker PID: {row.get('job_pid')} · Stop requested: {row.get('stop_requested')}")
                st.code(tail_text(live_log if coerce_text(live_log, '') else row.get('error_text'), limit=12000, fallback='(tanpa log)'), language='text')

def render_failures_tab(key_prefix: str = "failures"):
    st.markdown("### 🧪 Failure diagnostics")
    df = load_registry_df()
    if df.empty or "status" not in df.columns:
        st.info("Belum ada registry run.")
        return
    failed = df[df["status"] == "failed"].copy()
    if failed.empty:
        st.success("Belum ada run yang gagal tercatat di registry v4.0.")
        return
    display_cols = [c for c in ["run_id", "created_at", "experiment_name", "dataset", "workflow_mode", "experiment_type", "error_text"] if c in failed.columns]
    st.dataframe(failed[display_cols], use_container_width=True, hide_index=True)
    run_ids = failed["run_id"].dropna().tolist()
    selected = st.selectbox("Lihat traceback run gagal", options=run_ids, key=f"{key_prefix}_run")
    row = failed[failed["run_id"] == selected].iloc[0]
    st.code(coerce_text(row.get("error_text"), "(tanpa traceback)"), language="text")


def render_registry(key_prefix: str = "registry"):
    st.markdown("### 🗂️ Persistent experiment registry (SQLite)")
    df = load_registry_df()
    if df.empty:
        st.info("Belum ada eksperimen yang tersimpan pada registry persisten.")
        return
    c1, c2, c3 = st.columns(3)
    dataset_filter = c1.selectbox("Filter dataset", options=["All"] + sorted(df["dataset"].dropna().unique().tolist()), key=f"{key_prefix}_dataset_filter")
    status_filter = c2.selectbox("Filter status", options=["All"] + sorted(df["status"].dropna().unique().tolist()), key=f"{key_prefix}_status_filter")
    workflow_filter = c3.selectbox("Filter workflow", options=["All"] + sorted(df["workflow_mode"].dropna().unique().tolist()), key=f"{key_prefix}_workflow_filter")
    view = df.copy()
    if dataset_filter != "All":
        view = view[view["dataset"] == dataset_filter]
    if status_filter != "All":
        view = view[view["status"] == status_filter]
    if workflow_filter != "All":
        view = view[view["workflow_mode"] == workflow_filter]
    st.dataframe(view[[c for c in ["run_id", "created_at", "status", "experiment_name", "dataset", "classification_mode", "workflow_mode", "experiment_type", "best_model", "f1", "accuracy", "auc", "complexity_label"] if c in view.columns]], use_container_width=True, hide_index=True)
    csv_bytes = view.to_csv(index=False).encode("utf-8")
    xlsx_io = io.BytesIO()
    with pd.ExcelWriter(xlsx_io, engine="openpyxl") as writer:
        view.to_excel(writer, index=False, sheet_name="registry")
    st.download_button("⬇️ Registry CSV", data=csv_bytes, file_name="nids_registry_v3_4.csv", mime="text/csv", key=f"{key_prefix}_registry_csv")
    st.download_button("⬇️ Registry Excel", data=xlsx_io.getvalue(), file_name="nids_registry_v3_4.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=f"{key_prefix}_registry_excel")


def render_compare_board(latest: Optional[Dict], key_prefix: str = "compare"):
    st.markdown("### 📊 Compare board lintas sesi")
    df = load_registry_df()
    if df.empty:
        st.info("Belum ada run di registry untuk dibandingkan.")
        return
    metric = st.selectbox("Metric ranking", options=["f1", "accuracy", "auc"], key=f"{key_prefix}_metric")
    run_options = df["run_id"].tolist()[:50]
    selected = st.multiselect("Pilih run_id untuk dibandingkan", options=run_options, default=run_options[: min(5, len(run_options))], key=f"{key_prefix}_runs")
    compare_df = df[df["run_id"].isin(selected)].copy()
    if latest:
        summary_df = latest.get("result", {}).get("summary")
        best_model, best_f1, best_acc, best_auc = get_best_model_info(summary_df)
        latest_row = pd.DataFrame([{
            "run_id": latest.get("run_id"), "created_at": latest.get("timestamp"), "status": "session", "experiment_name": latest["config"].get("experiment_name"),
            "dataset": latest["config"].get("dataset"), "classification_mode": "Binary" if latest["config"].get("binary", True) else "Multiclass",
            "workflow_mode": latest["config"].get("workflow_mode"), "experiment_type": latest["config"].get("experiment_type"),
            "best_model": best_model, "f1": best_f1, "accuracy": best_acc, "auc": best_auc,
        }])
        compare_df = pd.concat([latest_row, compare_df], ignore_index=True)
    if compare_df.empty:
        st.info("Pilih minimal satu run dari registry.")
        return
    st.dataframe(compare_df[[c for c in ["run_id", "experiment_name", "dataset", "classification_mode", "workflow_mode", "experiment_type", "best_model", "f1", "accuracy", "auc"] if c in compare_df.columns]], use_container_width=True, hide_index=True)
    if metric in compare_df.columns:
        chart_df = compare_df[["run_id", metric]].dropna().set_index("run_id")
        if not chart_df.empty:
            st.bar_chart(chart_df)
    st.download_button("⬇️ Compare CSV", data=compare_df.to_csv(index=False).encode("utf-8"), file_name="compare_board_v3_1.csv", mime="text/csv", key=f"{key_prefix}_csv")


def render_leaderboard(key_prefix: str = "leader"):
    st.markdown("### 🏁 Leaderboard")
    df = load_registry_df()
    if df.empty:
        st.info("Belum ada run di registry.")
        return
    metric = st.selectbox("Leaderboard metric", options=["f1", "accuracy", "auc"], key=f"{key_prefix}_leader_metric")
    if metric not in df.columns:
        st.info("Metric belum tersedia di registry.")
        return
    view = df.dropna(subset=[metric]).sort_values(metric, ascending=False).head(10)
    for _, row in view.iterrows():
        st.markdown(
            f"<div class='leader-row'><strong>{row.get('experiment_name')}</strong> · {row.get('dataset')} · {row.get('classification_mode')}<br><span class='small-muted'>run_id={row.get('run_id')} · best_model={row.get('best_model')} · {metric}={float(row.get(metric)):.4f}</span></div>",
            unsafe_allow_html=True,
        )



def collect_preflight_issues(config: Dict) -> Tuple[List[str], List[str]]:
    issues: List[str] = []
    warnings: List[str] = []

    ml_active, dl_active = active_model_groups(config)
    requested_dl = config.get("requested_dl_models", []) or []
    tf_health = config.get("tf_health") or {}
    tf_ok = bool(tf_health.get("ok"))

    learning_paradigm = config.get("learning_paradigm", "supervised")
    if learning_paradigm in {"unsupervised", "reinforcement"} and not bool(config.get("binary", True)):
        issues.append(
            f"{LEARNING_PARADIGM_LABELS.get(learning_paradigm, learning_paradigm)} saat ini didukung untuk mode binary / anomaly detection. Pilih Binary classification."
        )
    if learning_paradigm == "reinforcement":
        warnings.append("Reinforcement Learning backend pada versi ini bersifat experimental adaptive-threshold policy, bukan DRL agent penuh.")

    if requested_dl and not tf_ok:
        dl_msg = coerce_text(tf_health.get("message"), fallback="TensorFlow/DL runtime belum siap.")
        if config.get("auto_disable_dl"):
            warnings.append(f"TensorFlow/DL runtime belum siap. Dashboard akan fallback ke ML-only untuk run ini. Detail: {dl_msg}")
        else:
            issues.append(f"TensorFlow/DL runtime belum siap sementara model DL dipilih. Aktifkan auto-fallback atau perbaiki environment DL. Detail: {dl_msg}")

    if (len(ml_active) + len(dl_active)) == 0:
        issues.append("Pilih minimal satu expert model ML atau DL sebelum menjalankan pipeline.")

    path_error = validate_input_paths(config.get("dataset", ""), config.get("paths", {}))
    if path_error:
        issues.append(path_error)

    pipeline_path = Path(str(config.get("pipeline_path", "")))
    if not pipeline_path.exists():
        issues.append(f"Pipeline file tidak ditemukan: {pipeline_path}")

    if config.get("execution_mode") == "foreground" and config.get("complexity_score", 0) >= 120:
        warnings.append("Eksperimen ini berat. Pertimbangkan mode background agar UI lebih stabil.")

    if config.get("workflow_mode") == "publication" and config.get("cv_folds", 2) < 5:
        warnings.append("Mode publication biasanya lebih kuat dengan CV >= 5.")

    if config.get("dataset") == "cicids" and config.get("sample_frac", 1.0) < 1.0 and config.get("workflow_mode") == "publication":
        warnings.append("Publication mode pada CICIDS lebih kuat bila sample fraction = 1.0 atau dijelaskan eksplisit sebagai exploratory/publication-lite.")
    if config.get("validation_method") == "nested_cv" and int(config.get("cv_folds", 2)) < 3:
        warnings.append("Nested CV biasanya lebih stabil dengan outer folds >= 3.")
    if config.get("feature_method") in ("chi2", "rfe", "mrmr"):
        warnings.append("Metode feature selection yang dipilih dapat menambah waktu komputasi. Pastikan ini konsisten dengan workflow yang diinginkan.")
    if config.get("optimize_ml"):
        warnings.append("Optimize ML aktif: runtime dapat meningkat karena RandomizedSearchCV pada TRAIN split.")
    if config.get("optimize_dl"):
        warnings.append("Optimize DL bersifat heuristic/lightweight, bukan AutoML penuh.")
    if config.get("use_hybrid") and (not ml_active or not dl_active):
        warnings.append("Hybrid aktif tetapi komposisi model belum lengkap. Runtime akan fallback ke non-hybrid agar run tidak gagal.")
    if config.get("experiment_type") in {"hybrid_fusion", "copula_bn"} and (not ml_active or not dl_active):
        issues.append("Experiment type hybrid/copula membutuhkan minimal satu model ML dan satu model DL aktif.")

    return issues, warnings

def apply_template_defaults(dataset: str, preset: str, workflow_mode: str, experiment_type: str) -> Dict:
    cfg = dict(PRESET_DEFAULTS[preset])
    if workflow_mode == "publication":
        cfg["cv_folds"] = max(cfg.get("cv_folds", 2), 5)
        cfg["sample_frac"] = 1.0
        cfg["use_shap"] = True
    elif workflow_mode == "exploratory":
        cfg["cv_folds"] = min(cfg.get("cv_folds", 2), 2)
        if dataset == "cicids":
            cfg["sample_frac"] = min(cfg.get("sample_frac", 0.3), 0.30)
    cfg.update(get_template_overrides(experiment_type))
    if experiment_type == "copula_bn":
        cfg["use_bn"] = True
    return cfg


def get_dataset_final_preset_overlay(dataset: str, preset_name: str) -> Dict:
    overlay = dict((DATASET_FINAL_PRESETS.get(dataset, {}) or {}).get(preset_name, {}) or {})
    overlay.pop("description", None)
    return overlay


def get_final_study_overlay(preset_name: str) -> Dict:
    return copy.deepcopy(FINAL_STUDY_PRESETS.get(preset_name, {}))

def get_publication_validation_overlay(preset_name: str) -> Dict:
    return copy.deepcopy(PUBLICATION_VALIDATION_PRESETS.get(preset_name, {}))

def merge_overlays(*items: Dict) -> Dict:
    merged = {}
    for d in items:
        if isinstance(d, dict):
            merged.update({k: v for k, v in d.items() if k != "description"})
    return merged


def get_dataset_final_preset_description(dataset: str, preset_name: str) -> str:
    return str((((DATASET_FINAL_PRESETS.get(dataset, {}) or {}).get(preset_name, {}) or {}).get("description", "")))


def build_sidebar() -> Tuple[bool, Dict]:
    st.sidebar.title("⚙️ Experiment Builder v4.6.1.1")
    tf_health = get_tensorflow_health()
    theme_choice = st.sidebar.selectbox("Theme", options=list(THEME_OPTIONS.keys()), index=list(THEME_OPTIONS.keys()).index(st.session_state.get("theme_choice", "cyber")), format_func=lambda x: THEME_OPTIONS[x]["label"], key="theme_choice_v341")
    st.session_state["theme_choice"] = theme_choice
    pipeline_path = st.sidebar.text_input("Path file pipeline (.py)", value=str(resolve_default_pipeline_path()), help="Arahkan ke file pipeline Python yang ingin dibungkus oleh GUI Streamlit ini.")
    with st.sidebar.expander("DL runtime health", expanded=False):
        if tf_health.get("ok"):
            st.success(coerce_text(tf_health.get("message"), "TensorFlow siap."))
        else:
            st.warning("TensorFlow / plugin DL belum siap. Anda masih bisa menjalankan mode ML-only.")
            st.caption(coerce_text(tf_health.get("message"), "Runtime DL belum siap."))
            tb = coerce_text(tf_health.get("traceback"))
            if tb:
                st.code(tail_text(tb, limit=4000, fallback="(tanpa traceback)"), language="text")

    st.sidebar.markdown("### 1) Research workflow")
    workflow_mode = st.sidebar.selectbox("Workflow mode", options=list(WORKFLOW_MODES.keys()), format_func=lambda x: WORKFLOW_MODES[x], index=0)
    experiment_type = st.sidebar.selectbox("Experiment type", options=list(EXPERIMENT_TYPES.keys()), format_func=lambda x: EXPERIMENT_TYPES[x], index=3)
    preset = st.sidebar.selectbox("Preset eksperimen", options=PRESET_OPTIONS, index=1)
    execution_mode = st.sidebar.radio("Execution mode", options=["background", "foreground"], format_func=lambda x: "Background queue" if x == "background" else "Foreground (blocking)", horizontal=False)
    auto_refresh_jobs = st.sidebar.checkbox("Auto-refresh jobs", value=st.session_state.get("jobs_auto_refresh", True))
    st.session_state["jobs_auto_refresh"] = auto_refresh_jobs
    auto_refresh_sec = st.sidebar.selectbox("Auto-refresh interval (sec)", options=[5, 10, 15, 30, 60], index=1)
    st.session_state["jobs_auto_refresh_interval"] = int(auto_refresh_sec)
    dataset = st.sidebar.selectbox("Dataset", options=list(DATASET_LABELS.keys()), format_func=lambda x: DATASET_LABELS[x])
    dataset_final_preset = st.sidebar.selectbox("Dataset final preset", options=list(DATASET_FINAL_PRESETS.get(dataset, {"None": {}}).keys()), index=0)
    study_preset = st.sidebar.selectbox("Final study preset", options=list(FINAL_STUDY_PRESETS.keys()), index=0)
    validation_preset = st.sidebar.selectbox("Publication validation preset", options=list(PUBLICATION_VALIDATION_PRESETS.keys()), index=0)
    dataset_final_preset_desc = get_dataset_final_preset_description(dataset, dataset_final_preset)
    if dataset_final_preset_desc:
        st.sidebar.caption(f"Preset dataset: **{dataset_final_preset}** — {dataset_final_preset_desc}")
    study_preset_desc = FINAL_STUDY_PRESETS.get(study_preset, {}).get("description")
    if study_preset_desc and study_preset != "Custom":
        st.sidebar.caption(f"Final study: **{study_preset}** — {study_preset_desc}")
    validation_preset_desc = PUBLICATION_VALIDATION_PRESETS.get(validation_preset, {}).get("description")
    if validation_preset_desc and validation_preset != "Custom":
        st.sidebar.caption(f"Validation: **{validation_preset}** — {validation_preset_desc}")
    dataset_overlay = merge_overlays(
        get_dataset_final_preset_overlay(dataset, dataset_final_preset),
        get_final_study_overlay(study_preset),
        get_publication_validation_overlay(validation_preset),
    )

    st.sidebar.markdown("### 2) Classification & preprocessing")
    binary = st.sidebar.radio("Mode klasifikasi", options=[True, False], index=0 if bool(dataset_overlay.get("binary", True)) else 1, format_func=lambda x: "Binary" if x else "Multiclass", horizontal=True)
    reco_scaler = SCALER_RECOMMENDATION[dataset]["recommended"]
    scaler_choice = st.sidebar.selectbox("Scaling", options=["auto", "robust", "standard", "minmax"], index=["auto", "robust", "standard", "minmax"].index(dataset_overlay.get("scaler", "auto") if dataset_overlay.get("scaler", "auto") in ["auto", "robust", "standard", "minmax"] else "auto"), format_func=lambda x: "Auto recommended" if x == "auto" else f"{x} — {SCALER_LABELS.get(x, x)}")
    resolved_scaler = reco_scaler if scaler_choice == "auto" else scaler_choice
    st.sidebar.caption(f"Rekomendasi: **{reco_scaler}** — {SCALER_RECOMMENDATION[dataset]['reason']}")
    balance = st.sidebar.selectbox("Balancing", options=list(BALANCE_LABELS.keys()), index=list(BALANCE_LABELS.keys()).index(dataset_overlay.get("balance", list(BALANCE_LABELS.keys())[0])) if dataset_overlay.get("balance", list(BALANCE_LABELS.keys())[0]) in list(BALANCE_LABELS.keys()) else 0, format_func=lambda x: f"{x} — {BALANCE_LABELS[x]}")
    copula_choice = st.sidebar.selectbox("Copula", options=["auto"] + list(COPULA_LABELS.keys()), index=(["auto"] + list(COPULA_LABELS.keys())).index(dataset_overlay.get("copula_family", "auto") if dataset_overlay.get("copula_family", "auto") in (["auto"] + list(COPULA_LABELS.keys())) else "auto"), format_func=lambda x: "Auto default" if x == "auto" else f"{x} — {COPULA_LABELS[x]}")
    resolved_copula = COPULA_DEFAULT.get(dataset, "gaussian") if copula_choice == "auto" else copula_choice

    st.sidebar.markdown("### 3) Data source")
    prepared_input = render_upload_inputs(dataset)
    render_path_status(prepared_input.paths, dataset)

    template_cfg = apply_template_defaults(dataset, preset, workflow_mode, experiment_type)
    template_cfg.update({k: v for k, v in dataset_overlay.items() if k in MODEL_FLAGS or k in {"cv_folds", "top_k", "sample_frac", "use_bn", "use_shap"}})
    if st.sidebar.button("Apply template to model/library settings", use_container_width=True, key="apply_template_v32"):
        apply_template_to_state(template_cfg)
        st.rerun()
    for key in MODEL_FLAGS:
        st.session_state.setdefault(f"model_{key}", template_cfg.get(key, False))

    st.sidebar.markdown("### 4) Core parameters")
    top_k = st.sidebar.slider("Top-k features", min_value=10, max_value=100, value=int(dataset_overlay.get("top_k", template_cfg.get("top_k", 40))), step=5)
    cv_default = int(dataset_overlay.get("cv_folds", template_cfg.get("cv_folds", 2)))
    cv_folds = st.sidebar.slider("CV folds", min_value=2, max_value=10, value=cv_default, step=1)
    sample_frac = 1.0
    cicids_split = "random"
    if dataset == "cicids":
        sample_frac = st.sidebar.slider("Sample fraction (CICIDS)", min_value=0.05, max_value=1.0, value=float(dataset_overlay.get("sample_frac", template_cfg.get("sample_frac", 0.30))), step=0.05)
        cicids_split = st.sidebar.selectbox("Split mode (CICIDS)", options=["random", "temporal"], index=1 if workflow_mode == "publication" else 0)
    seed = st.sidebar.number_input("Seed / run control", min_value=0, max_value=999999, value=42, step=1)
    validation_method = st.sidebar.selectbox("Validation method", options=list(VALIDATION_METHOD_LABELS.keys()), format_func=lambda x: VALIDATION_METHOD_LABELS[x], index=list(VALIDATION_METHOD_LABELS.keys()).index(dataset_overlay.get("validation_method", "holdout") if dataset_overlay.get("validation_method", "holdout") in VALIDATION_METHOD_LABELS else "holdout"))
    validation_repeats = 3
    bootstrap_rounds = 30
    nested_inner_folds = 3
    if validation_method == "repeated_kfold":
        validation_repeats = st.sidebar.slider("Repeated K-Fold repeats", min_value=2, max_value=10, value=3, step=1)
    elif validation_method == "bootstrapping":
        bootstrap_rounds = st.sidebar.slider("Bootstrap rounds", min_value=10, max_value=200, value=30, step=5)
    elif validation_method == "nested_cv":
        nested_inner_folds = st.sidebar.slider("Nested CV inner folds", min_value=2, max_value=10, value=3, step=1)
    feature_method = st.sidebar.selectbox("Feature selection", options=list(FEATURE_METHOD_LABELS.keys()), format_func=lambda x: FEATURE_METHOD_LABELS[x], index=list(FEATURE_METHOD_LABELS.keys()).index(dataset_overlay.get("feature_method", "mutual_info") if dataset_overlay.get("feature_method", "mutual_info") in FEATURE_METHOD_LABELS else "mutual_info"))
    optimize_ml = st.sidebar.checkbox("Optimize ML models", value=(workflow_mode == "publication"), help="Lightweight RandomizedSearchCV pada TRAIN split saja.")
    optimize_dl = st.sidebar.checkbox("Optimize DL heuristically", value=False, help="Heuristic tuning untuk epochs/batch size, bukan AutoML penuh.")
    use_ensemble = st.sidebar.checkbox("Enable ensemble methods", value=experiment_type in ("ml_ensemble", "dl_ensemble", "hybrid_fusion", "copula_bn"))
    ensemble_method = st.sidebar.selectbox("Ensemble method", options=list(ENSEMBLE_METHOD_LABELS.keys()), format_func=lambda x: ENSEMBLE_METHOD_LABELS[x], index=list(ENSEMBLE_METHOD_LABELS.keys()).index("weighted_soft_vote"))
    use_hybrid = st.sidebar.checkbox("Enable hybrid methods", value=experiment_type in ("hybrid_fusion", "copula_bn"))
    hybrid_method = st.sidebar.selectbox("Hybrid method", options=list(HYBRID_METHOD_LABELS.keys()), format_func=lambda x: HYBRID_METHOD_LABELS[x], index=list(HYBRID_METHOD_LABELS.keys()).index("weighted_ml_dl"))

    st.sidebar.markdown("### 5) Learning paradigm & model taxonomy")
    learning_paradigm = st.sidebar.selectbox(
        "Learning paradigm",
        options=list(LEARNING_PARADIGM_LABELS.keys()),
        format_func=lambda x: LEARNING_PARADIGM_LABELS[x],
        index=0,
        help="Level utama pengelompokan model. Pada v4.5.1, backend supervised, semi-supervised, unsupervised, dan RL experimental sudah bisa dijalankan.",
    )
    supervised_track = "both"
    auto_disable_dl = True

    if learning_paradigm == "supervised":
        supervised_track = st.sidebar.radio(
            "Supervised track",
            options=list(SUPERVISED_TRACK_LABELS.keys()),
            format_func=lambda x: SUPERVISED_TRACK_LABELS[x],
            index=2,
            help="Pilih sub-struktur menu untuk kelompok model pada paradigma supervised.",
        )

        show_ml = supervised_track in {"ml", "both", "hybrid"}
        show_dl = supervised_track in {"dl", "both", "hybrid"}

        if show_ml:
            with st.sidebar.expander("Supervised → Model ML", expanded=True):
                c1, c2 = st.columns(2)
                for i, (key, short, note) in enumerate(ML_MODEL_SPECS):
                    with [c1, c2][i % 2]:
                        st.checkbox(short, key=f"model_{key}", help=note)

        if show_dl:
            with st.sidebar.expander("Supervised → Model DL", expanded=True):
                c1, c2 = st.columns(2)
                for i, (key, short, note) in enumerate(DL_MODEL_SPECS):
                    with [c1, c2][i % 2]:
                        st.checkbox(short, key=f"model_{key}", help=note)
                auto_disable_dl = st.checkbox("Auto-fallback ke ML-only bila TensorFlow tidak siap", value=True, key="auto_disable_dl_v45", help="Jika TensorFlow/plugin DL bermasalah, semua pilihan DL akan dinonaktifkan pada konfigurasi run sehingga pipeline tetap bisa berjalan dengan model ML.")
                if not tf_health.get("ok"):
                    st.warning("Runtime DL tidak sehat saat ini. Model DL dapat dinonaktifkan otomatis atau Anda sinkronkan manual.")
                    if st.button("Sinkronkan sekarang: nonaktifkan semua model DL", use_container_width=True, key="disable_all_dl_now_v45"):
                        for _k, _, _ in DL_MODEL_SPECS:
                            st.session_state[f"model_{_k}"] = False
                        st.rerun()
                else:
                    st.success("Runtime DL terdeteksi siap dipakai.")
        if supervised_track == "hybrid":
            st.sidebar.caption("Track hybrid/fusion idealnya mengaktifkan minimal satu model ML dan satu model DL, lalu dipadukan melalui ensemble / hybrid method / copula.")

    elif learning_paradigm == "semi_supervised":
        with st.sidebar.expander("Semi-Supervised → Model ML", expanded=True):
            c1, c2 = st.columns(2)
            for i, (key, short, note) in enumerate(SEMI_ML_MODEL_SPECS):
                with [c1, c2][i % 2]:
                    st.checkbox(short, key=f"model_{key}", help=note)
        with st.sidebar.expander("Semi-Supervised → Model DL", expanded=True):
            c1, c2 = st.columns(2)
            for i, (key, short, note) in enumerate(SEMI_DL_MODEL_SPECS):
                with [c1, c2][i % 2]:
                    st.checkbox(short, key=f"model_{key}", help=note)
            auto_disable_dl = st.checkbox("Auto-fallback: nonaktifkan DL semi-supervised bila TensorFlow tidak siap", value=True, key="auto_disable_dl_semi_v45")
            if not tf_health.get("ok"):
                st.warning("Runtime DL belum sehat; model semi-supervised DL dapat dinonaktifkan otomatis.")
    elif learning_paradigm == "unsupervised":
        with st.sidebar.expander("Unsupervised → Model ML", expanded=True):
            c1, c2 = st.columns(2)
            for i, (key, short, note) in enumerate(UNSUP_ML_MODEL_SPECS):
                with [c1, c2][i % 2]:
                    st.checkbox(short, key=f"model_{key}", help=note)
        with st.sidebar.expander("Unsupervised → Model DL", expanded=True):
            for key, short, note in UNSUP_DL_MODEL_SPECS:
                st.checkbox(short, key=f"model_{key}", help=note)
            auto_disable_dl = st.checkbox("Auto-fallback: nonaktifkan DL anomaly bila TensorFlow tidak siap", value=True, key="auto_disable_dl_unsup_v45")
            if not tf_health.get("ok"):
                st.warning("Runtime DL belum sehat; autoencoder anomaly dapat dinonaktifkan otomatis.")
        st.sidebar.caption("Paradigma unsupervised / anomaly detection saat ini ditujukan untuk mode binary.")
    else:
        with st.sidebar.expander("Reinforcement Learning → Model ML", expanded=True):
            for key, short, note in RL_ML_MODEL_SPECS:
                st.checkbox(short, key=f"model_{key}", help=note)
        with st.sidebar.expander("Reinforcement Learning → Model DL (roadmap)", expanded=False):
            for label in ROADMAP_MODEL_GROUPS.get("reinforcement", {}).get("DL", []):
                st.checkbox(label, value=False, disabled=True, key=f"roadmap_rl_{label}")
        st.sidebar.info("RL pada v4.5.1 bersifat experimental: adaptive threshold Q-policy di atas anomaly scores, cocok sebagai research scaffold.")
    st.sidebar.markdown("### 6) Advanced hyperparameters")
    hp_defaults = HYPERPARAM_DEFAULTS.copy()
    with st.sidebar.expander("Common DL training controls", expanded=False):
        dl_epochs = st.number_input("DL epochs", min_value=1, max_value=500, value=int(hp_defaults["dl_epochs"]), step=1, key="dl_epochs_v32")
        dl_batch_size = st.selectbox("DL batch size", options=[32, 64, 128, 256, 512], index=[32, 64, 128, 256, 512].index(int(hp_defaults["dl_batch_size"])), key="dl_batch_v32")
        dl_learning_rate = st.number_input("DL learning rate (informational/global)", min_value=0.0001, max_value=0.1, value=0.0010, step=0.0005, format="%.4f", key="dl_lr_v32")
        dl_dropout_hint = st.slider("DL dropout hint", min_value=0.0, max_value=0.7, value=0.30, step=0.05, key="dl_dropout_v32")
    with st.sidebar.expander("ML model hyperparameters", expanded=False):
        lr_c = st.number_input("LR: C", min_value=0.001, max_value=1000.0, value=float(hp_defaults["lr_c"]), step=0.1, key="lr_c_v32")
        lr_solver = st.selectbox("LR: solver", options=["lbfgs", "liblinear", "saga", "newton-cg"], index=0, key="lr_solver_v32")
        lr_penalty = st.selectbox("LR: penalty", options=["l2", "l1", "elasticnet", None], index=0, key="lr_penalty_v32")
        rf_n_estimators = st.number_input("RF: n_estimators", min_value=10, max_value=2000, value=int(hp_defaults["rf_n_estimators"]), step=10, key="rf_n_estimators_v32")
        rf_max_depth = st.number_input("RF: max_depth (0=None)", min_value=0, max_value=200, value=int(hp_defaults["rf_max_depth"]), step=1, key="rf_max_depth_v32")
        rf_min_samples_leaf = st.number_input("RF: min_samples_leaf", min_value=1, max_value=50, value=2, step=1, key="rf_min_leaf_v32")
        svm_c = st.number_input("SVM: C", min_value=0.001, max_value=1000.0, value=float(hp_defaults["svm_c"]), step=0.1, key="svm_c_v32")
        svm_kernel = st.selectbox("SVM: kernel", options=["rbf", "linear", "poly", "sigmoid"], index=0, key="svm_kernel_v32")
        svm_gamma = st.selectbox("SVM: gamma", options=["scale", "auto"], index=0, key="svm_gamma_v32")
        xgb_n_estimators = st.number_input("XGB: n_estimators", min_value=10, max_value=3000, value=int(hp_defaults["xgb_n_estimators"]), step=10, key="xgb_n_estimators_v32")
        xgb_max_depth = st.number_input("XGB: max_depth", min_value=1, max_value=32, value=int(hp_defaults["xgb_max_depth"]), step=1, key="xgb_max_depth_v32")
        xgb_learning_rate = st.number_input("XGB: learning_rate", min_value=0.001, max_value=1.0, value=float(hp_defaults["xgb_learning_rate"]), step=0.01, format="%.3f", key="xgb_learning_rate_v32")
        xgb_subsample = st.slider("XGB: subsample", min_value=0.3, max_value=1.0, value=0.8, step=0.05, key="xgb_subsample_v32")
        xgb_colsample = st.slider("XGB: colsample_bytree", min_value=0.3, max_value=1.0, value=0.8, step=0.05, key="xgb_colsample_v32")
        dt_max_depth = st.number_input("DT: max_depth", min_value=1, max_value=200, value=int(hp_defaults["dt_max_depth"]), step=1, key="dt_max_depth_v32")
        dt_min_samples_leaf = st.number_input("DT: min_samples_leaf", min_value=1, max_value=50, value=2, step=1, key="dt_min_leaf_v32")
        dt_criterion = st.selectbox("DT: criterion", options=["gini", "entropy", "log_loss"], index=0, key="dt_criterion_v32")

    st.sidebar.markdown("### 6b) Paradigm-specific controls")
    semi_label_fraction = st.sidebar.slider("Semi-supervised labeled fraction", min_value=0.05, max_value=0.95, value=0.30, step=0.05, help="Porsi data TRAIN yang diperlakukan sebagai berlabel pada metode semi-supervised.")
    pseudo_label_threshold = st.sidebar.slider("Pseudo-label confidence threshold", min_value=0.50, max_value=0.99, value=0.90, step=0.01)
    anomaly_contamination = st.sidebar.slider("Anomaly contamination / outlier rate", min_value=0.01, max_value=0.50, value=0.10, step=0.01)
    rl_episodes = st.sidebar.slider("RL episodes (experimental)", min_value=10, max_value=200, value=40, step=5)

    st.sidebar.markdown("### 7) Auxiliary analytics")
    aux1, aux2 = st.sidebar.columns(2)
    use_bn = aux1.checkbox("Bayesian Net", value=bool(template_cfg.get("use_bn", True)), key="use_bn_v32")
    use_shap = aux2.checkbox("SHAP", value=bool(template_cfg.get("use_shap", False)), key="use_shap_v32")

    st.sidebar.markdown("### 8) Data quality & EDA")
    generate_eda = st.sidebar.checkbox("Generate EDA figures", value=True, help="Jika dimatikan, pipeline akan melewati tahap EDA berat dan hanya menjalankan training/evaluasi utama.")
    drop_duplicates_stage = st.sidebar.checkbox("Drop duplicate rows in hygiene stage", value=True)
    drop_constant_stage = st.sidebar.checkbox("Drop constant features in hygiene stage", value=True)

    st.sidebar.markdown("### 9) Protocol lock, notes & formal report metadata")
    experiment_name = st.sidebar.text_input("Nama eksperimen", value=f"{dataset}_{workflow_mode}_{experiment_type}")
    notes = st.sidebar.text_area("Catatan eksperimen", value="", height=90, placeholder="Mis. binary baseline robust + gaussian")
    protocol_lock = st.sidebar.checkbox("Saya mengonfirmasi rancangan eksperimen ini sudah final untuk dijalankan", value=False)
    report_template = st.sidebar.selectbox("Formal report template", options=["formal_dissertation", "formal_journal", "thesis", "paper", "concise"], index=0)
    report_title = st.sidebar.text_input("Formal report title", value=f"Hybrid NIDS Experiment on {DATASET_LABELS.get(dataset, dataset)}")
    report_author = st.sidebar.text_input("Author", value="Researcher")
    report_affiliation = st.sidebar.text_input("Affiliation", value="Graduate Research Program")

    requested_model_flags = {k: st.session_state.get(f"model_{k}", False) for k in MODEL_FLAGS}
    if learning_paradigm == "supervised":
        # disable all non-supervised families
        for _spec in (SEMI_ML_MODEL_SPECS + SEMI_DL_MODEL_SPECS + UNSUP_ML_MODEL_SPECS + UNSUP_DL_MODEL_SPECS + RL_ML_MODEL_SPECS + RL_DL_MODEL_SPECS):
            requested_model_flags[_spec[0]] = False
        if supervised_track == "ml":
            for _k, _, _ in DL_MODEL_SPECS:
                requested_model_flags[_k] = False
        elif supervised_track == "dl":
            for _k, _, _ in ML_MODEL_SPECS:
                requested_model_flags[_k] = False
    elif learning_paradigm == "semi_supervised":
        for _spec in (ML_MODEL_SPECS + DL_MODEL_SPECS + UNSUP_ML_MODEL_SPECS + UNSUP_DL_MODEL_SPECS + RL_ML_MODEL_SPECS + RL_DL_MODEL_SPECS):
            requested_model_flags[_spec[0]] = False
    elif learning_paradigm == "unsupervised":
        for _spec in (ML_MODEL_SPECS + DL_MODEL_SPECS + SEMI_ML_MODEL_SPECS + SEMI_DL_MODEL_SPECS + RL_ML_MODEL_SPECS + RL_DL_MODEL_SPECS):
            requested_model_flags[_spec[0]] = False
    else:
        for _spec in (ML_MODEL_SPECS + DL_MODEL_SPECS + SEMI_ML_MODEL_SPECS + SEMI_DL_MODEL_SPECS + UNSUP_ML_MODEL_SPECS + UNSUP_DL_MODEL_SPECS):
            requested_model_flags[_spec[0]] = False

    all_dl_specs = DL_MODEL_SPECS + SEMI_DL_MODEL_SPECS + UNSUP_DL_MODEL_SPECS + RL_DL_MODEL_SPECS
    requested_dl_models = [label for key, label, _ in all_dl_specs if requested_model_flags.get(key)]
    if (not tf_health.get("ok")) and auto_disable_dl:
        for _k, _, _ in all_dl_specs:
            requested_model_flags[_k] = False
    effective_disabled_dl = [label for key, label, _ in all_dl_specs if (not requested_model_flags.get(key)) and st.session_state.get(f"model_{key}", False)]

    config = {
        "pipeline_path": pipeline_path,
        "dataset": dataset,
        "paths": prepared_input.paths,
        "learning_paradigm": learning_paradigm,
        "supervised_track": supervised_track,
        "paradigm_backend_supported": True,
        "binary": binary,
        "scaler": resolved_scaler,
        "balance": balance,
        "copula_family": resolved_copula,
        "top_k": top_k,
        "cv_folds": cv_folds,
        "validation_method": validation_method,
        "validation_repeats": int(validation_repeats),
        "bootstrap_rounds": int(bootstrap_rounds),
        "nested_inner_folds": int(nested_inner_folds),
        "feature_method": feature_method,
        "optimize_ml": bool(optimize_ml),
        "optimize_dl": bool(optimize_dl),
        "use_ensemble": bool(use_ensemble),
        "ensemble_method": ensemble_method,
        "use_hybrid": bool(use_hybrid),
        "hybrid_method": hybrid_method,
        "sample_frac": sample_frac,
        "cicids_split": cicids_split,
        **requested_model_flags,
        "use_bn": use_bn,
        "use_shap": use_shap,
        "generate_eda": bool(generate_eda),
        "drop_duplicates_stage": bool(drop_duplicates_stage),
        "drop_constant_stage": bool(drop_constant_stage),
        "preset": preset,
        "dataset_final_preset": dataset_final_preset,
        "study_preset": study_preset,
        "validation_preset": validation_preset,
        "resolved_scaler": resolved_scaler,
        "resolved_copula": resolved_copula,
        "reco_scaler": reco_scaler,
        "experiment_name": experiment_name,
        "input_mode": prepared_input.mode,
        "saved_input_files": [str(p) for p in prepared_input.saved_files],
        "notes": notes,
        "theme": theme_choice,
        "workflow_mode": workflow_mode,
        "experiment_type": experiment_type,
        "seed": int(seed),
        "protocol_lock": protocol_lock,
        "execution_mode": execution_mode,
        "tf_health": tf_health,
        "auto_disable_dl": bool(auto_disable_dl),
        "requested_dl_models": requested_dl_models,
        "forced_disabled_dl_models": effective_disabled_dl,
        "training_params": {
            "epochs": int(dl_epochs),
            "batch_size": int(dl_batch_size),
            "learning_rate_hint": float(dl_learning_rate),
            "dropout_hint": float(dl_dropout_hint),
            "seed": int(seed),
        },
        "model_params": {
            "lr": {"C": float(lr_c), "solver": lr_solver, "penalty": lr_penalty},
            "rf": {"n_estimators": int(rf_n_estimators), "max_depth": (None if int(rf_max_depth) == 0 else int(rf_max_depth)), "min_samples_leaf": int(rf_min_samples_leaf)},
            "svm": {"C": float(svm_c), "kernel": svm_kernel, "gamma": svm_gamma},
            "xgb": {"n_estimators": int(xgb_n_estimators), "max_depth": int(xgb_max_depth), "learning_rate": float(xgb_learning_rate), "subsample": float(xgb_subsample), "colsample_bytree": float(xgb_colsample)},
            "dt": {"max_depth": int(dt_max_depth), "min_samples_leaf": int(dt_min_samples_leaf), "criterion": dt_criterion},
        },
    }
    score, complexity_label, complexity_note = estimate_complexity(config)
    config["complexity_score"] = score
    config["complexity_label"] = complexity_label
    config["complexity_note"] = complexity_note
    ml_active, dl_active = active_model_groups(config)
    config["active_ml_models"] = ml_active
    config["active_dl_models"] = dl_active
    config["tf_runtime_ok"] = bool(tf_health.get("ok"))

    preflight_issues, preflight_warnings = collect_preflight_issues(config)
    config["preflight_issues"] = preflight_issues
    config["preflight_warnings"] = preflight_warnings

    st.sidebar.markdown("---")
    st.sidebar.markdown(f"**Complexity:** {complexity_label} ({score})")
    st.sidebar.progress(min(score / 170.0, 1.0))
    st.sidebar.caption(complexity_note)
    st.sidebar.caption(
        f"Paradigm: {LEARNING_PARADIGM_LABELS.get(config.get('learning_paradigm', 'supervised'))} · "
        f"ML aktif: {len(ml_active)} · DL aktif: {len(dl_active)}"
    )

    with st.sidebar.expander("Preflight checks", expanded=bool(preflight_issues or preflight_warnings)):
        if preflight_issues:
            for msg in preflight_issues:
                st.error(msg)
        else:
            st.success("Tidak ada blocker utama pada konfigurasi saat ini.")
        for msg in preflight_warnings:
            st.warning(msg)

    disabled = (not protocol_lock) or bool(preflight_issues)
    button_label = "🚀 Submit background job" if execution_mode == "background" else "🚀 Run pipeline now"
    run_clicked = st.sidebar.button(button_label, use_container_width=True, type="primary", disabled=disabled)
    if not protocol_lock:
        st.sidebar.caption("Aktifkan protocol lock untuk menjalankan pipeline.")
    elif preflight_issues:
        st.sidebar.caption("Perbaiki preflight issues terlebih dahulu sebelum menjalankan pipeline.")
    elif execution_mode == "background":
        st.sidebar.caption("Run akan masuk queue dan dikerjakan di background worker.")
    return run_clicked, config


def main():
    st.set_page_config(page_title="Hybrid NIDS Research Studio v4.6.1 Final Study", page_icon="🛡️", layout="wide")
    init_state()
    reconcile_background_jobs()
    started_run_id = start_next_queued_job()
    inject_css(st.session_state.get("theme_choice", "cyber"))
    render_login_gate()

    run_clicked, config = build_sidebar()
    inject_css(config.get("theme", st.session_state.get("theme_choice", "cyber")))
    render_account_box()

    if st.session_state.get("result_bundle") is None:
        latest_loaded = load_latest_success_bundle()
        if latest_loaded is not None:
            st.session_state["result_bundle"] = latest_loaded

    validation_error = validate_input_paths(config["dataset"], config["paths"])
    audit = None
    if not validation_error:
        try:
            audit = audit_dataset_cached(config["dataset"], json.dumps(config["paths"], sort_keys=True))
        except Exception as exc:
            audit = {"overall": {}, "files": [], "error": str(exc)}
    guards = compute_guardrails(config, audit, validation_error)
    readiness_score, readiness_label, readiness_notes = compute_readiness(config, validation_error, audit, guards)

    latest = st.session_state.get("result_bundle")
    render_hero(config, history_count=len(st.session_state["history"]), readiness_score=readiness_score, readiness_label=readiness_label)
    if started_run_id:
        st.info(f"Background worker dimulai untuk job: {started_run_id}")
    render_overview_cards(latest, config, guards, audit)
    render_config_panel(config, readiness_score, readiness_label, readiness_notes)
    render_tf_health_panel(config.get("tf_health") or {}, compact=False)
    startup_report = run_startup_self_test(str(config.get("pipeline_path", "")), config.get("tf_health") or {})
    render_startup_self_test_panel(startup_report, key_prefix="startup_main")
    with st.sidebar.expander("Startup self-check", expanded=False):
        render_startup_self_test_panel(startup_report, key_prefix="startup_sidebar")

    tabs = st.tabs(["Dashboard", "Jobs", "Audit", "Preprocessing", "Study Design", "Results", "Figures", "Logs", "Diagnostics", "Registry", "Compare", "Report", "Reproducibility", "Leaderboard"])

    with tabs[0]:
        st.markdown("### 📍 Dashboard overview")
        st.markdown(f"<div class='glass'><div class='metric-label'>Dataset note</div><div class='small-muted'>{SPLIT_PROTOCOLS[config['dataset']]}</div></div>", unsafe_allow_html=True)
        if latest:
            render_summary(latest, key_prefix=f"dashboard_summary_{latest.get('run_id', 'latest')}")
        else:
            st.info("Belum ada hasil run pada session ini.")

    with tabs[1]:
        render_jobs_tab(key_prefix="jobs_tab")

    with tabs[2]:
        render_dataset_audit(audit)
        render_guardrails(guards)

    with tabs[3]:
        render_prepreprocessing_report(latest)
        if latest:
            render_analytics_gallery(latest, key_prefix=f"prep_gallery_{latest.get('run_id', 'latest')}")
        else:
            st.info("Jalankan pipeline untuk melihat audit dan grafik analytics.")

    with tabs[4]:
        render_study_design(config, latest, audit, guards)

    with tabs[5]:
        if latest:
            render_summary(latest, key_prefix=f"results_summary_{latest.get('run_id', 'latest')}")
            package_exists = latest.get("artifact_zip_path") and Path(latest["artifact_zip_path"]).exists()
            if package_exists:
                st.download_button("⬇️ Download academic package (.zip)", data=Path(latest["artifact_zip_path"]).read_bytes(), file_name=Path(latest["artifact_zip_path"]).name, mime="application/zip", key=f"results_pkg_{latest.get('run_id')}_v43")
        else:
            st.info("Belum ada hasil run.")

    with tabs[6]:
        if latest:
            render_figures(latest, key_prefix=f"figs_{latest.get('run_id', 'latest')}")
        else:
            st.info("Belum ada figure untuk ditampilkan.")

    with tabs[7]:
        if latest:
            render_logs(latest, key_prefix=f"logs_{latest.get('run_id', 'latest')}")
        else:
            st.info("Belum ada log untuk ditampilkan.")

    with tabs[8]:
        render_failures_tab(key_prefix="diag_tab")

    with tabs[9]:
        render_registry(key_prefix="registry_tab")

    with tabs[10]:
        render_compare_board(latest, key_prefix="compare_tab")

    with tabs[11]:
        render_report_tab(latest, config, audit, guards, key_prefix="report_tab")

    with tabs[12]:
        render_reproducibility_tab(config, latest, audit, guards, key_prefix="repro_tab")

    with tabs[13]:
        render_leaderboard(key_prefix="leader_tab")

    if validation_error:
        st.error(validation_error)

    if run_clicked:
        if config.get("execution_mode") == "background":
            run_id = enqueue_background_job(config, audit, guards)
            st.session_state["run_counter"] += 1
            start_next_queued_job()
            st.success(f"Job {run_id} masuk ke background queue.")
            st.rerun()
        else:
            run_id = make_run_id(config)
            st.session_state["run_counter"] += 1
            try:
                module, module_path = load_pipeline_module(config["pipeline_path"])
            except Exception as exc:
                st.error(f"Gagal memuat pipeline: {exc}")
                st.stop()
            progress = st.progress(0.02)
            status_box = st.empty()
            log_box = st.empty()
            status_box.info("Menyiapkan registry dan menjalankan pipeline...")
            db_insert_running(run_id, config, audit, guards)
            try:
                progress.progress(0.08)
                result, live_logs, elapsed = run_pipeline(module, config, log_box)
                progress.progress(0.82)
                fig_paths, log_paths = collect_output_files(module_path.parent, config["dataset"])
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                latest = {
                    "run_id": run_id,
                    "timestamp": timestamp,
                    "result": result,
                    "logs": live_logs,
                    "elapsed_sec": elapsed,
                    "fig_paths": fig_paths,
                    "log_paths": log_paths,
                    "config": config,
                    "pipeline_path": str(module_path),
                }
                manifest = make_methodology_manifest(latest, audit=audit, guards=guards)
                report_md = make_academic_report_markdown(latest, config=config, audit=audit, guards=guards, style="thesis")
                out_dir, zip_path = save_artifacts(run_id, latest, audit, guards, manifest, report_md)
                latest["artifact_zip_path"] = str(zip_path)
                latest["artifact_root"] = str(out_dir)
                st.session_state["result_bundle"] = latest
                st.session_state["history"].append(make_history_row(run_id, latest, manifest))
                db_complete_run(run_id, status="success", config=config, manifest=manifest, audit=audit, guards=guards, result_bundle=latest, artifact_zip_path=zip_path, output_root=out_dir)
                progress.progress(1.0)
                status_box.success(f"Pipeline selesai dalam {elapsed:.1f} detik. Artifact package siap diunduh.")
                st.rerun()
            except Exception as exc:
                err = traceback.format_exc()
                db_complete_run(run_id, status="failed", config=config, manifest=None, audit=audit, guards=guards, result_bundle=None, artifact_zip_path=None, output_root=None, error_text=err)
                progress.progress(1.0)
                status_box.error(f"Pipeline gagal: {exc}")
                if "proba_dict kosong" in str(exc) or "Tidak ada probabilitas model" in str(exc):
                    st.info("Petunjuk: ini biasanya berarti tidak ada model yang aktif, semua model gagal build/train, atau semua output prediksi gagal. Coba aktifkan minimal satu model ML sederhana seperti LR / RF / DT dan periksa dependency.")
                log_box.code(err, language="text")

    st.markdown("<div class='footer-note'>v4.6.1 menambahkan preset eksperimen final per dataset, startup self-test yang lebih detail, serta formal academic reporting dengan metadata judul, penulis, dan afiliasi untuk ekspor Markdown, DOCX, dan PDF.</div>", unsafe_allow_html=True)


DATASET_NOTES = {
    "nslkdd": "Official train/test split dipertahankan; validation diambil dari official train saja.",
    "unsw": "Official training/testing split dipertahankan; validation diambil dari training set saja.",
    "insdn": "Beberapa sumber traffic digabung terlebih dahulu lalu dibagi train/val/test.",
    "cicids": "Banyak file flow harian digabung dulu, lalu split random atau temporal.",
}


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "--worker":
        raise SystemExit(worker_main(sys.argv[2]))
    main()